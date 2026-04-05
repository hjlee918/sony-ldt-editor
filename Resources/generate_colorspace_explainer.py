from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ─── Page margins ───
section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(0.85)
section.bottom_margin = Inches(0.85)

# ─── Palette ───
BLACK      = RGBColor(0x1A, 0x1A, 0x1A)
DARK_GRAY  = RGBColor(0x44, 0x44, 0x44)
MID_GRAY   = RGBColor(0x77, 0x77, 0x77)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
SONY_BLUE  = RGBColor(0x00, 0x3B, 0x8E)
ACCENT     = RGBColor(0x00, 0x6D, 0xC6)
AMBER      = RGBColor(0xB8, 0x6C, 0x00)
GREEN_DARK = RGBColor(0x1A, 0x6B, 0x2E)
RED_DARK   = RGBColor(0x99, 0x1A, 0x1A)

def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'),   kwargs[edge].get('val','single'))
            tag.set(qn('w:sz'),    kwargs[edge].get('sz','4'))
            tag.set(qn('w:space'),'0')
            tag.set(qn('w:color'), kwargs[edge].get('color','auto'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def remove_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'none')
        tblBorders.append(tag)
    tblPr.append(tblBorders)

def add_rule(color_hex='CCCCCC', thickness=4, space_before=0, space_after=8):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(thickness))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    return p

def add_heading(text, level=1, color=SONY_BLUE, size=None, space_before=18,
                space_after=6, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = color
    run.font.size = Pt(size or {1:20,2:14,3:12,4:11}.get(level,11))
    run.font.name = 'Calibri'
    return p

def add_body(text, color=DARK_GRAY, size=10.5, italic=False,
             space_before=2, space_after=6, bold=False):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name      = 'Calibri'
    run.font.size      = Pt(size)
    run.font.color.rgb = color
    run.italic         = italic
    run.bold           = bold
    return p

def add_mixed(parts, space_before=2, space_after=6):
    """parts = list of (text, bold, color, italic)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    for text, bold, color, italic in parts:
        run = p.add_run(text)
        run.font.name      = 'Calibri'
        run.font.size      = Pt(10.5)
        run.font.color.rgb = color
        run.bold           = bold
        run.italic         = italic
    return p

def add_callout(text, fill_hex='EEF4FF', border_hex='006DC6',
                label=None, label_color=ACCENT, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(10)
    p.paragraph_format.left_indent  = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    shade_paragraph(p, fill_hex)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for edge in ('top','left','bottom','right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   'single')
        tag.set(qn('w:sz'),    '12' if edge == 'left' else '4')
        tag.set(qn('w:space'), '4')
        tag.set(qn('w:color'), border_hex)
        pBdr.append(tag)
    pPr.append(pBdr)
    if label:
        rl = p.add_run(label + '   ')
        rl.bold = True
        rl.font.name  = 'Calibri'
        rl.font.size  = Pt(size)
        rl.font.color.rgb = label_color
    rb = p.add_run(text)
    rb.font.name      = 'Calibri'
    rb.font.size      = Pt(size)
    rb.font.color.rgb = DARK_GRAY
    return p

def add_mono_block(lines, fill_hex='F4F4F4'):
    """Monospaced diagram block."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.5)
        shade_paragraph(p, fill_hex)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GRAY
    # spacer after block
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after  = Pt(8)

def add_step_table(rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(table)

    # header
    hdr = table.rows[0]
    for i, h in enumerate(['Step', 'Control', 'Why']):
        cell = hdr.cells[i]
        shade_cell(cell, '003B8E')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p   = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(h.upper())
        run.bold = True
        run.font.name      = 'Calibri'
        run.font.size      = Pt(9)
        run.font.color.rgb = WHITE

    for ri, (step, ctrl, why) in enumerate(rows):
        tr   = table.rows[ri + 1]
        fill = 'F7F9FF' if ri % 2 == 1 else 'FFFFFF'
        data = [step, ctrl, why]
        for ci, text in enumerate(data):
            cell = tr.cells[ci]
            shade_cell(cell, fill)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell, bottom={'val':'single','sz':'2','color':'E0E0E0'})
            p   = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            run = p.add_run(text)
            run.font.name      = 'Calibri'
            run.font.size      = Pt(9.5)
            run.font.color.rgb = BLACK
            run.bold = (ci == 0)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

# ═══════════════════════════════════════════════════════════
#  COVER
# ═══════════════════════════════════════════════════════════
p_cover = doc.add_paragraph()
p_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cover.paragraph_format.space_before = Pt(4)
p_cover.paragraph_format.space_after  = Pt(2)
shade_paragraph(p_cover, '003B8E')
rc = p_cover.add_run('Color Space Gamut Adjustment')
rc.font.name      = 'Calibri'
rc.font.size      = Pt(20)
rc.font.color.rgb = WHITE
rc.bold           = True

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after  = Pt(4)
shade_paragraph(p_sub, '003B8E')
rs = p_sub.add_run('Understanding Cyan–Red & Magenta–Green Sliders  ·  VPL-VW385ES')
rs.font.name      = 'Calibri'
rs.font.size      = Pt(10)
rs.font.color.rgb = RGBColor(0xAA, 0xC8, 0xFF)

add_rule('1A6FD8', 8, space_before=0, space_after=10)

# ═══════════════════════════════════════════════════════════
#  SECTION 1 — THE CIE CHROMATICITY DIAGRAM
# ═══════════════════════════════════════════════════════════
add_heading('1  The CIE Chromaticity Diagram', level=1, space_before=6)
add_body(
    'To understand these sliders, you first need a mental model of how colors are mapped '
    'in the CIE chromaticity diagram — the universal map of all visible colors.'
)
add_body(
    'Every color a human eye can see has a precise (x, y) coordinate on this map. '
    'The map is shaped like a horseshoe, with saturated spectral colors along the curved '
    'edge and white near the center. A display\'s color gamut is represented as a triangle '
    'whose three corners are its Red, Green, and Blue primaries. Every color the display '
    'can produce sits somewhere inside that triangle.'
)

add_mono_block([
    '                  G (Green primary)',
    '                 / \\',
    '                /   \\',
    '               /  W  \\      W = White point (e.g. D65)',
    '              /       \\',
    '             B─────────R',
    '     (Blue primary)   (Red primary)',
])

add_body(
    'The shape of the triangle defines which colors are available. The position of the '
    'white point W inside the triangle determines the color temperature. Both of these '
    'are set by the Color Space mode you choose (BT.709, BT.2020, etc.).'
)

# ═══════════════════════════════════════════════════════════
#  SECTION 2 — WHAT THE SLIDERS DO
# ═══════════════════════════════════════════════════════════
add_heading('2  What Cyan–Red and Magenta–Green Actually Do', level=1)
add_body(
    'These two sliders move the entire gamut triangle as a single unit across the '
    'chromaticity map — all three primaries (R, G, B) shift together in the same '
    'direction, by the same amount. The shape and size of the triangle does not change.'
)

add_heading('The Two Axes', level=2, color=ACCENT, space_before=10, space_after=4)
add_body(
    'The chromaticity map has two dimensions, and these sliders each control one of them:'
)

# Two-column explanation boxes
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
remove_table_borders(table)

for ci, (title, col_hex, border_hex, direction, effect) in enumerate([
    ('Cyan – Red', 'FFF0F0', 'C43030',
     'Negative (–) → toward Cyan\nPositive (+) → toward Red',
     'Shifts the whole triangle horizontally. Pushing positive makes all colors warmer/redder overall; pushing negative makes them cooler/more cyan.'),
    ('Magenta – Green', 'F0FFF4', '1A7A42',
     'Negative (–) → toward Magenta\nPositive (+) → toward Green',
     'Shifts the whole triangle vertically. Pushing positive gives all colors a slight greenish cast; pushing negative gives a magenta/purple cast.'),
]):
    cell = table.rows[0].cells[ci]
    shade_cell(cell, col_hex)
    set_cell_border(cell,
        top={'val':'single','sz':'8','color':border_hex},
        left={'val':'single','sz':'4','color':border_hex},
        right={'val':'single','sz':'4','color':border_hex},
        bottom={'val':'single','sz':'4','color':border_hex})
    cell.width = Inches(3.0)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(title)
    r.bold = True
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(
        int(border_hex[0:2],16),
        int(border_hex[2:4],16),
        int(border_hex[4:6],16)
    )

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after  = Pt(2)
    r2 = p2.add_run(direction)
    r2.font.name      = 'Calibri'
    r2.font.size      = Pt(9)
    r2.font.color.rgb = DARK_GRAY
    r2.italic         = True

    p3 = cell.add_paragraph()
    p3.paragraph_format.space_before = Pt(4)
    p3.paragraph_format.space_after  = Pt(6)
    r3 = p3.add_run(effect)
    r3.font.name      = 'Calibri'
    r3.font.size      = Pt(9.5)
    r3.font.color.rgb = DARK_GRAY

sp = doc.add_paragraph()
sp.paragraph_format.space_after = Pt(10)

add_body(
    'Together, these two sliders give you full 2D control over where the gamut triangle '
    'sits on the chromaticity map — you can nudge it in any direction by combining both.'
)

# Stamp analogy callout
add_callout(
    'Think of the gamut triangle as a rubber stamp. The Cyan–Red and Magenta–Green '
    'sliders control where you press it down on the page — left/right and up/down. '
    'The shape of the stamp never changes, only its position.',
    fill_hex='FFFBEA', border_hex='B86C00', label='ANALOGY', label_color=AMBER
)

# Diagram showing the shift
add_heading('Visualising the Shift', level=2, color=ACCENT, space_before=10, space_after=4)
add_body(
    'Below: the dashed triangle is the original BT.709 gamut position. '
    'The solid triangle shows what happens after applying a positive Cyan–Red adjustment '
    '(the whole gamut shifts rightward — toward red — on the chromaticity map).'
)

add_mono_block([
    '              G′                   G = original Green primary',
    '             /·\\                   G′= shifted Green primary',
    '            / · \\    shift →',
    '           /  · ·\\',
    '          B′· · · R′              All three corners move',
    '         ·B·······R               together by the same amount',
])

# ═══════════════════════════════════════════════════════════
#  SECTION 3 — EFFECT ON CMS CALIBRATION
# ═══════════════════════════════════════════════════════════
add_heading('3  Effect on CMS Calibration', level=1)
add_body(
    'This is where these sliders have the most significant practical consequence for calibration work.'
)

add_heading('What CMS Does', level=2, color=ACCENT, space_before=10, space_after=4)
add_body(
    'The Color Correction CMS (6-color Hue / Saturation / Brightness adjustments for '
    'Red, Yellow, Green, Cyan, Blue, Magenta) works by fine-tuning where each individual '
    'primary and secondary color sits relative to its target position. A calibrator '
    'measures each color with a colorimeter, finds the error, and dials in corrections '
    'until each color lands on its target coordinates.'
)

add_heading('What Happens When You Move the Gamut After CMS', level=2,
            color=RED_DARK, space_before=10, space_after=4)
add_body(
    'If you adjust Cyan–Red or Magenta–Green after CMS calibration has been performed, '
    'the entire gamut shifts underneath the CMS corrections. Every color moves away from '
    'the position the CMS corrections were pointing at. The corrections are now wrong '
    'for every single color simultaneously.'
)

add_callout(
    'Analogy: you carefully parked each car (R, G, B, C, M, Y) in its exact designated '
    'space. Then someone moved the entire car park. All cars are still in their same '
    'spaces relative to each other — but every car is now in the wrong space. '
    'The CMS must be redone from scratch.',
    fill_hex='FFF0F0', border_hex='C43030', label='WARNING', label_color=RED_DARK
)

add_body(
    'In short: yes — adjusting these sliders after CMS work invalidates all per-color '
    'corrections. The global gamut position is the foundation that CMS is built on top of.'
)

# ═══════════════════════════════════════════════════════════
#  SECTION 4 — CORRECT CALIBRATION ORDER
# ═══════════════════════════════════════════════════════════
add_heading('4  Correct Calibration Order', level=1)
add_body(
    'Because the Color Space gamut position is the foundation for everything else, '
    'it must be set first and locked before any downstream calibration is performed.'
)

add_step_table(
    rows=[
        ('1', 'Color Space\n(+ Cyan–Red / Magenta–Green)',
         'Sets the gamut foundation — where the triangle sits on the chromaticity map. Lock this first.'),
        ('2', 'Color Temperature\n(Gain R/G/B · Bias R/G/B)',
         'Sets the white point and black point within the chosen gamut.'),
        ('3', 'Gamma Correction\n(or upload custom .ldt curve)',
         'Sets the tonal response — how brightness maps from input to output.'),
        ('4', 'Color Correction CMS\n(Hue / Sat / Brightness per color)',
         'Fine-tunes individual primaries and secondaries last, after the gamut foundation is locked.'),
    ],
    col_widths=[0.45, 2.1, 3.85]
)

sp = doc.add_paragraph()
sp.paragraph_format.space_after = Pt(6)

add_callout(
    'Return to step 1 at any point and all subsequent steps must be repeated. '
    'Steps 2, 3, and 4 are always downstream of the Color Space gamut position.',
    fill_hex='EEF4FF', border_hex='006DC6', label='RULE', label_color=ACCENT
)

# ═══════════════════════════════════════════════════════════
#  SECTION 5 — PRACTICAL USE CASES
# ═══════════════════════════════════════════════════════════
add_heading('5  When to Actually Use These Sliders', level=1)
add_body(
    'In a well-functioning projector measured against its target color space, these '
    'sliders should ideally sit at 0. They exist to compensate for real-world variations:'
)

table2 = doc.add_table(rows=3, cols=2)
table2.alignment = WD_TABLE_ALIGNMENT.LEFT
remove_table_borders(table2)

use_cases = [
    ('Systematic gamut error',
     'If measurement shows the projector\'s gamut is consistently shifted in one direction '
     'vs. the BT.709 targets (e.g. slightly redward across all colors), a small Cyan–Red '
     'correction fixes the whole gamut at once — more efficient than correcting every '
     'individual color in CMS.'),
    ('Unit-to-unit matching',
     'Two projectors of the same model can have slight gamut position differences from '
     'lamp and panel manufacturing tolerances. This global trim brings them into alignment '
     'before per-color CMS work.'),
    ('Room/environment compensation',
     'Ambient light with a strong color cast (e.g. warm tungsten room lighting) can be '
     'partially compensated here at the gamut level before fine CMS work.'),
]
for ri, (title, desc) in enumerate(use_cases):
    fill = 'F7F9FF' if ri % 2 == 1 else 'FFFFFF'
    c0, c1 = table2.rows[ri].cells[0], table2.rows[ri].cells[1]
    for cell in (c0, c1):
        shade_cell(cell, fill)
        set_cell_border(cell, bottom={'val':'single','sz':'2','color':'E0E0E0'})

    p0 = c0.paragraphs[0]
    p0.paragraph_format.space_before = Pt(5)
    p0.paragraph_format.space_after  = Pt(5)
    r0 = p0.add_run(title)
    r0.bold = True
    r0.font.name      = 'Calibri'
    r0.font.size      = Pt(9.5)
    r0.font.color.rgb = SONY_BLUE

    p1 = c1.paragraphs[0]
    p1.paragraph_format.space_before = Pt(5)
    p1.paragraph_format.space_after  = Pt(5)
    r1 = p1.add_run(desc)
    r1.font.name      = 'Calibri'
    r1.font.size      = Pt(9.5)
    r1.font.color.rgb = DARK_GRAY

for row in table2.rows:
    row.cells[0].width = Inches(1.8)
    row.cells[1].width = Inches(4.6)

sp2 = doc.add_paragraph()
sp2.paragraph_format.space_after = Pt(8)

# ─── Footer ───
add_rule('CCCCCC', 4, space_before=12, space_after=4)
pf = doc.add_paragraph()
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = pf.add_run(
    'Sony VPL-VW385ES  ·  Color Space Gamut Adjustment Reference  ·  '
    'Verified from OSD screenshot and operating instructions (47254410M)'
)
rf.font.name      = 'Calibri'
rf.font.size      = Pt(8)
rf.font.color.rgb = MID_GRAY

# ─── Save ───
out = '/Users/johnlee/Downloads/sony-ldt-editor/Resources/ColorSpace_Gamut_Adjustment_Explainer.docx'
doc.save(out)
print(f'Saved: {out}')
