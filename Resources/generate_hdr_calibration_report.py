"""
Generates: HDR Calibration Procedure for Sony VPL-VW385ES
Tools: CalMAN Business 5.12, i1Display Pro Plus, ArgyllPro v2,
       PGenerator 1.6 (RPi4), Sony LDT Editor, Sony ImageDirector
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Palette ───
BLACK      = RGBColor(0x1A,0x1A,0x1A)
DARK_GRAY  = RGBColor(0x44,0x44,0x44)
MID_GRAY   = RGBColor(0x77,0x77,0x77)
WHITE      = RGBColor(0xFF,0xFF,0xFF)
SONY_BLUE  = RGBColor(0x00,0x3B,0x8E)
ACCENT     = RGBColor(0x00,0x6D,0xC6)
AMBER      = RGBColor(0xB8,0x6C,0x00)
RED_DARK   = RGBColor(0x99,0x1A,0x1A)
GREEN_DARK = RGBColor(0x1A,0x6B,0x2E)
PURPLE     = RGBColor(0x5B,0x2D,0x8E)
TEAL       = RGBColor(0x0E,0x7A,0x72)

# Hex versions
H_BLUE  = '003B8E'; H_ACCENT= '006DC6'; H_AMBER = 'B86C00'
H_RED   = '991A1A'; H_GREEN = '1A6B2E'; H_PURPLE= '5B2D8E'
H_TEAL  = '0E7A72'; H_GRAY  = '555555'

# Phase colors
PHASE_COLORS = {
    0: ('1A1A2E','E8E8FF'),  # Setup - dark navy
    1: ('1A4A1A','E8FFE8'),  # Baseline - dark green
    2: ('4A1A00','FFE8D0'),  # White Balance - dark orange
    3: ('00254A','D0E8FF'),  # Gamma/EOTF - dark blue
    4: ('2D004A','EED0FF'),  # CMS - dark purple
    5: ('004A40','D0FFF8'),  # Verification - dark teal
    6: ('1A1A1A','F0F0F0'),  # Save/Document - dark gray
}

def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_width=Inches(8.5); s.page_height=Inches(11)
    s.left_margin=Inches(0.9); s.right_margin=Inches(0.9)
    s.top_margin=Inches(0.8); s.bottom_margin=Inches(0.8)
    return doc

def shade_para(p, hex_color):
    pPr=p._p.get_or_add_pPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'),hex_color); pPr.append(shd)

def shade_cell(cell, hex_color):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'),hex_color); tcPr.append(shd)

def border_cell(cell, **kw):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        if edge in kw:
            t=OxmlElement(f'w:{edge}')
            t.set(qn('w:val'),kw[edge].get('val','single'))
            t.set(qn('w:sz'),kw[edge].get('sz','4'))
            t.set(qn('w:space'),'0')
            t.set(qn('w:color'),kw[edge].get('color','auto'))
            b.append(t)
    tcPr.append(b)

def no_borders(table):
    tbl=table._tbl; tblPr=tbl.find(qn('w:tblPr'))
    if tblPr is None: tblPr=OxmlElement('w:tblPr'); tbl.insert(0,tblPr)
    tb=OxmlElement('w:tblBorders')
    for e in ('top','left','bottom','right','insideH','insideV'):
        t=OxmlElement(f'w:{e}'); t.set(qn('w:val'),'none'); tb.append(t)
    tblPr.append(tb)

def hrule(doc, color='CCCCCC', sz=4, sb=4, sa=8):
    p=doc.add_paragraph(); pPr=p._p.get_or_add_pPr()
    pBdr=OxmlElement('w:pBdr'); bot=OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),str(sz))
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),color)
    pBdr.append(bot); pPr.append(pBdr)
    p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)

def heading(doc, text, level=1, color=SONY_BLUE, size=None,
            sb=14, sa=6, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT):
    p=doc.add_paragraph(); p.alignment=align
    p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(text); r.bold=bold; r.font.name='Calibri'
    r.font.color.rgb=color
    r.font.size=Pt(size or {1:17,2:13,3:11,4:10}.get(level,10))
    return p

def body(doc, text, color=DARK_GRAY, size=10.5, italic=False,
         sb=2, sa=6, bold=False):
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(text); r.font.name='Calibri'; r.font.size=Pt(size)
    r.font.color.rgb=color; r.italic=italic; r.bold=bold
    return p

def callout(doc, text, fill='EEF4FF', border='006DC6',
            label=None, lcolor=None, size=10):
    lc=lcolor or ACCENT
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(10)
    p.paragraph_format.left_indent=Inches(0.15)
    p.paragraph_format.right_indent=Inches(0.15)
    shade_para(p, fill); pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
    for edge in ('top','left','bottom','right'):
        t=OxmlElement(f'w:{edge}')
        t.set(qn('w:val'),'single')
        t.set(qn('w:sz'),'16' if edge=='left' else '4')
        t.set(qn('w:space'),'4'); t.set(qn('w:color'),border); pBdr.append(t)
    pPr.append(pBdr)
    if label:
        rl=p.add_run(label+'   '); rl.bold=True
        rl.font.name='Calibri'; rl.font.size=Pt(size); rl.font.color.rgb=lc
    rb=p.add_run(text); rb.font.name='Calibri'
    rb.font.size=Pt(size); rb.font.color.rgb=DARK_GRAY

def mono_block(doc, lines, fill='F0F0F0', title=None):
    if title:
        pt=doc.add_paragraph()
        pt.paragraph_format.space_before=Pt(6); pt.paragraph_format.space_after=Pt(0)
        shade_para(pt,'2D2D2D')
        rt=pt.add_run('  '+title)
        rt.font.name='Calibri'; rt.font.size=Pt(8.5)
        rt.font.color.rgb=RGBColor(0xAA,0xCC,0xFF); rt.bold=True
    for line in lines:
        p=doc.add_paragraph()
        p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
        p.paragraph_format.left_indent=Inches(0.3)
        shade_para(p, '1E1E1E' if title else fill)
        r=p.add_run(line)
        r.font.name='Courier New'; r.font.size=Pt(9)
        r.font.color.rgb=RGBColor(0xD4,0xD4,0xD4) if title else DARK_GRAY
    sp=doc.add_paragraph()
    sp.paragraph_format.space_before=Pt(0); sp.paragraph_format.space_after=Pt(8)

def data_table(doc, headers, rows, col_widths=None,
               hdr_fill='003B8E', alt_fill='F5F7FF', note=None):
    t=doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment=WD_TABLE_ALIGNMENT.LEFT; no_borders(t)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; shade_cell(c, hdr_fill)
        c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
        p=c.paragraphs[0]
        p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(5)
        r=p.add_run(h); r.bold=True; r.font.name='Calibri'
        r.font.size=Pt(8.5); r.font.color.rgb=WHITE
    for ri,row_data in enumerate(rows):
        tr=t.rows[ri+1]; fill=alt_fill if ri%2==1 else 'FFFFFF'
        for ci,txt in enumerate(row_data):
            c=tr.cells[ci]; shade_cell(c,fill)
            c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            border_cell(c,bottom={'val':'single','sz':'2','color':'E0E0E0'})
            p=c.paragraphs[0]
            p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)
            parts=txt.split('|B|')
            for pi,part in enumerate(parts):
                r=p.add_run(part); r.font.name='Calibri'; r.font.size=Pt(9.5)
                r.font.color.rgb=BLACK; r.bold=(ci==0 or pi%2==1)
    if col_widths:
        for i,w in enumerate(col_widths):
            for row in t.rows: row.cells[i].width=Inches(w)
    if note:
        pn=doc.add_paragraph()
        pn.paragraph_format.space_before=Pt(2); pn.paragraph_format.space_after=Pt(8)
        rn=pn.add_run('Note: '+note); rn.font.name='Calibri'
        rn.font.size=Pt(8.5); rn.font.color.rgb=MID_GRAY; rn.italic=True
    else:
        doc.add_paragraph().paragraph_format.space_after=Pt(6)
    return t

def phase_banner(doc, phase_num, phase_name, subtitle=''):
    fill,_ = PHASE_COLORS.get(phase_num,('003B8E','FFFFFF'))
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(16); p.paragraph_format.space_after=Pt(2)
    shade_para(p, fill)
    badge=f'  PHASE {phase_num}  '
    r1=p.add_run(badge); r1.bold=True; r1.font.name='Calibri'
    r1.font.size=Pt(9); r1.font.color.rgb=WHITE
    r2=p.add_run('  '+phase_name.upper())
    r2.bold=True; r2.font.name='Calibri'
    r2.font.size=Pt(13); r2.font.color.rgb=WHITE
    if subtitle:
        p2=doc.add_paragraph()
        p2.paragraph_format.space_before=Pt(0); p2.paragraph_format.space_after=Pt(10)
        shade_para(p2, fill)
        r3=p2.add_run('  '+subtitle)
        r3.font.name='Calibri'; r3.font.size=Pt(9.5)
        r3.font.color.rgb=RGBColor(0xCC,0xCC,0xCC)
    else:
        doc.add_paragraph().paragraph_format.space_after=Pt(8)

def step_row(doc, num, text, detail=None, color=SONY_BLUE):
    t=doc.add_table(rows=1,cols=2); no_borders(t)
    c0=t.rows[0].cells[0]; c1=t.rows[0].cells[1]
    shade_cell(c0, H_BLUE); c0.width=Inches(0.38)
    border_cell(c0,top={'val':'single','sz':'2','color':'FFFFFF'},
                   bottom={'val':'single','sz':'2','color':'FFFFFF'})
    p0=c0.paragraphs[0]; p0.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before=Pt(5); p0.paragraph_format.space_after=Pt(5)
    r0=p0.add_run(str(num)); r0.bold=True; r0.font.name='Calibri'
    r0.font.size=Pt(11); r0.font.color.rgb=WHITE
    shade_cell(c1,'FAFAFA'); c1.width=Inches(6.12)
    border_cell(c1,bottom={'val':'single','sz':'2','color':'E0E0E0'})
    p1=c1.paragraphs[0]
    p1.paragraph_format.space_before=Pt(5); p1.paragraph_format.space_after=Pt(2 if detail else 5)
    r1=p1.add_run(text); r1.bold=True; r1.font.name='Calibri'
    r1.font.size=Pt(10.5); r1.font.color.rgb=BLACK
    if detail:
        p2=c1.add_paragraph()
        p2.paragraph_format.space_before=Pt(0); p2.paragraph_format.space_after=Pt(5)
        r2=p2.add_run(detail); r2.font.name='Calibri'
        r2.font.size=Pt(9.5); r2.font.color.rgb=DARK_GRAY
    doc.add_paragraph().paragraph_format.space_after=Pt(3)

def tolerance_table(doc):
    headers=['Metric','Excellent','Good','Acceptable','Action Required']
    rows=[
        ['White Point dE2000',    '< 1.0','1.0 – 2.0','2.0 – 3.0','> 3.0 — re-adjust Gain/Bias'],
        ['Primary dE2000 (R/G/B)','< 1.5','1.5 – 3.0','3.0 – 4.0','> 4.0 — re-adjust CMS'],
        ['Secondary dE2000 (C/M/Y)','< 2.0','2.0 – 3.5','3.5 – 5.0','> 5.0 — re-adjust CMS'],
        ['EOTF avg deviation',    '< 3%', '3–5%',     '5–8%',     '> 8% — re-upload PQ curve'],
        ['Gamma at 50% stimulus', '± 0.05','± 0.10',  '± 0.15',   '> ± 0.15 — re-check curve'],
        ['Luminance at 100% white','± 5%', '± 10%',   '± 15%',    '> ± 15% — re-check lamp/iris'],
        ['Black level (nit)',     '< 0.002','< 0.005','< 0.010',   '> 0.01 — check Advanced Iris'],
        ['Saturation sweep dE2000','< 2.0','2.0–3.5', '3.5–5.0',  '> 5.0 — note as known limitation'],
    ]
    t=doc.add_table(rows=1+len(rows),cols=5)
    t.alignment=WD_TABLE_ALIGNMENT.LEFT; no_borders(t)
    hdr_colors=['003B8E','1A6B2E','006DC6','B86C00','991A1A']
    for i,(h,hc) in enumerate(zip(headers,hdr_colors)):
        c=t.rows[0].cells[i]; shade_cell(c,hc)
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)
        r=p.add_run(h); r.bold=True; r.font.name='Calibri'
        r.font.size=Pt(8); r.font.color.rgb=WHITE
    row_fills=[
        ['FFFFFF','E8FFE8','D0E8FF','FFE8CC','FFD0D0'],
        ['F5F5F5','E0FFE0','C8E0FF','FFE0BB','FFC8C8'],
    ]
    for ri,row_data in enumerate(rows):
        tr=t.rows[ri+1]; fills=row_fills[ri%2]
        for ci,txt in enumerate(row_data):
            c=tr.cells[ci]; shade_cell(c,fills[ci])
            c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            border_cell(c,bottom={'val':'single','sz':'2','color':'E0E0E0'})
            p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if ci>0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
            r=p.add_run(txt); r.font.name='Calibri'; r.font.size=Pt(9); r.font.color.rgb=BLACK
            r.bold=(ci==0)
    widths=[2.0,1.0,1.0,1.0,1.5]
    for i,w in enumerate(widths):
        for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(6)

def footer(doc, text):
    hrule(doc,'CCCCCC',4,sb=14,sa=4)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.font.name='Calibri'
    r.font.size=Pt(8); r.font.color.rgb=MID_GRAY

# ═══════════════════════════════════════════════════════════
doc = new_doc()

# ─── COVER ───
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(0)
shade_para(p,'001830')
r=p.add_run('HDR Calibration Procedure')
r.font.name='Calibri'; r.font.size=Pt(22); r.font.color.rgb=WHITE; r.bold=True

p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before=Pt(0); p2.paragraph_format.space_after=Pt(0)
shade_para(p2,'001830')
r2=p2.add_run('Sony VPL-VW385ES  ·  HDR10 / Rec.2020 / PQ EOTF')
r2.font.name='Calibri'; r2.font.size=Pt(12); r2.font.color.rgb=RGBColor(0x88,0xBB,0xFF)

p3=doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before=Pt(4); p3.paragraph_format.space_after=Pt(6)
shade_para(p3,'001830')
r3=p3.add_run('CalMAN Business 5.12  ·  i1Display Pro Plus  ·  PGenerator 1.6 (RPi4)  ·  ArgyllPro ColorMeter v2')
r3.font.name='Calibri'; r3.font.size=Pt(9); r3.font.color.rgb=RGBColor(0x66,0x99,0xCC)

hrule(doc,'1A5FA0',10,sb=0,sa=8)

# ─── TOOL OVERVIEW CARDS ───
heading(doc,'Tool Overview',1,sb=4,sa=6)
body(doc,'Six tools work together in a defined sequence. Each has a specific role — none is interchangeable.',sb=2,sa=8)

t_tools=doc.add_table(rows=2,cols=3); t_tools.alignment=WD_TABLE_ALIGNMENT.LEFT; no_borders(t_tools)
tools=[
    ('PATTERN\nGENERATOR','PGenerator 1.6\n(Raspberry Pi 4)','Outputs HDR10 test patches via HDMI directly to projector. Controlled by CalMAN over Wi-Fi/LAN. Ted\'s calibration patterns built in.','0E4A7A'),
    ('COLORIMETER','X-Rite i1Display\nPro Plus (2000 nit)','Measures light from the projection screen. Capable of 2000 nit range — well above the ~200 nit peak of VW385ES.','5B2D8E'),
    ('CALIBRATION\nSOFTWARE','CalMAN Business\n5.12','Controls PGenerator, reads i1Display Pro, displays dE errors, guides CMS adjustments. Primary calibration engine.','1A6B2E'),
]
tools2=[
    ('CURVE\nGENERATOR','Sony LDT Editor\n(Web App)','Generates BT.1886 and PQ tone-mapping curves as .ldt files for upload. Runs in any browser.','B86C00'),
    ('UPLOAD TOOL','Sony\nImageDirector','Uploads .ldt curve files from PC to the projector\'s Gamma 7–10 slots via LAN connection.','003B8E'),
    ('VERIFICATION','ArgyllPro\nColorMeter v2','Independent second-opinion measurements. Saturation sweep, ICC profiling, dE reporting — separate from CalMAN.','0E7A72'),
]
for ri,tool_row in enumerate([tools,tools2]):
    for ci,(role,name,desc,hexc) in enumerate(tool_row):
        c0=t_tools.rows[ri].cells[ci]
        shade_cell(c0,'F8F8F8')
        border_cell(c0,
            top={'val':'single','sz':'12','color':hexc},
            left={'val':'single','sz':'4','color':'E0E0E0'},
            right={'val':'single','sz':'4','color':'E0E0E0'},
            bottom={'val':'single','sz':'4','color':'E0E0E0'})
        c0.width=Inches(2.2)
        pRole=c0.paragraphs[0]
        pRole.paragraph_format.space_before=Pt(6); pRole.paragraph_format.space_after=Pt(2)
        rRole=pRole.add_run(role); rRole.bold=True; rRole.font.name='Calibri'
        rRole.font.size=Pt(7.5)
        rRole.font.color.rgb=RGBColor(int(hexc[0:2],16),int(hexc[2:4],16),int(hexc[4:6],16))
        pName=c0.add_paragraph()
        pName.paragraph_format.space_before=Pt(0); pName.paragraph_format.space_after=Pt(4)
        rName=pName.add_run(name); rName.bold=True; rName.font.name='Calibri'
        rName.font.size=Pt(10); rName.font.color.rgb=BLACK
        pDesc=c0.add_paragraph()
        pDesc.paragraph_format.space_before=Pt(0); pDesc.paragraph_format.space_after=Pt(8)
        rDesc=pDesc.add_run(desc); rDesc.font.name='Calibri'
        rDesc.font.size=Pt(8.5); rDesc.font.color.rgb=DARK_GRAY

doc.add_paragraph().paragraph_format.space_after=Pt(8)

# ─── SYSTEM DIAGRAM ───
heading(doc,'System Connection Diagram',1,sa=6)
mono_block(doc,[
    '                    ┌─────────────────────────────────────────────────────┐',
    '                    │           LOCAL AREA NETWORK (Wi-Fi / LAN)          │',
    '                    └───────────────┬─────────────────┬───────────────────┘',
    '                                    │                 │',
    '         ┌──────────────────────────▼──┐         ┌───▼──────────────────────┐',
    '         │  CALIBRATION PC / MacBook   │         │  SONY VPL-VW385ES        │',
    '         │  ┌─────────────────────┐   │         │  (LAN port connected)    │',
    '         │  │  CalMAN Business    │   │         └───────────────────────────┘',
    '         │  │  5.12               │   │                    ▲',
    '         │  └──────┬──────────────┘   │                    │ HDMI',
    '         │         │ USB              │         ┌──────────┴───────────────┐',
    '         │  ┌──────▼──────────────┐   │         │  RASPBERRY PI 4          │',
    '         │  │  i1Display Pro Plus │   │◄────────│  PGenerator 1.6          │',
    '         │  │  (on screen)        │   │  LAN    │  Ted\'s HDR Patterns      │',
    '         │  └─────────────────────┘   │         └──────────────────────────┘',
    '         │                            │',
    '         │  ┌─────────────────────┐   │',
    '         │  │  ArgyllPro v2       │   │   ← Also connects to i1Display Pro',
    '         │  │  (verification)     │   │     via USB for independent reads',
    '         │  └─────────────────────┘   │',
    '         │                            │',
    '         │  ┌─────────────────────┐   │',
    '         │  │  LDT Editor (browser│   │   ← Generates .ldt files',
    '         │  │  + ImageDirector    │   │     ImageDirector uploads via LAN',
    '         │  └─────────────────────┘   │',
    '         └────────────────────────────┘',
],'F4F6F8')

# ─── TARGETS ───
heading(doc,'HDR Calibration Targets  —  VPL-VW385ES',1)
data_table(doc,
    headers=['Parameter','Target Value','Notes'],
    rows=[
        ['Color Space Container','BT.2020','HDR10 signal container'],
        ['Display Color Gamut','DCI-P3 D65','Most HDR10 content mastered within P3 — calibrate CMS to P3 primaries'],
        ['White Point','D65  (x=0.3127, y=0.3290)','HDR mastering white point'],
        ['EOTF','PQ / ST.2084','Uploaded as custom .ldt curve to Gamma 8 slot'],
        ['Peak Luminance','~200 nit','VW385ES measured peak — used as PQ tone-map target'],
        ['Black Level','< 0.002 nit','Projector with Advanced Iris: Full'],
        ['HDR Setting (projector)','HDR10 (manual)','Do NOT use Auto during calibration'],
        ['Calib. Preset','Cinema Film 1','Store all HDR calibration settings here'],
        ['Gamma Slot','Gamma 8','PQ curve uploaded via ImageDirector'],
        ['Color Temp Slot','Custom 3 (D65)','Adjust Gain/Bias R/G/B here'],
    ],
    col_widths=[2.1,2.0,2.5],hdr_fill='001830')

# ─── DCI-P3 TARGETS BOX ───
mono_block(doc,[
    '  DCI-P3 D65 Primary Chromaticity Targets:',
    '',
    '  Red      x = 0.680   y = 0.320',
    '  Green    x = 0.265   y = 0.690',
    '  Blue     x = 0.150   y = 0.060',
    '  White    x = 0.3127  y = 0.3290  (D65)',
    '',
    '  Enter these as targets in CalMAN → Color Space settings.',
],'1E1E2E',title='CalMAN Target Coordinates — Enter These in CalMAN Color Space Settings')

# ══════════════════════════════════════════════════════════════
#  PHASE 0 — ENVIRONMENT & SETUP
# ══════════════════════════════════════════════════════════════
phase_banner(doc,0,'Environment & Setup','Before powering anything on — room, equipment, and software preparation')

heading(doc,'Room Conditions',2,color=ACCENT,sb=6,sa=4)
data_table(doc,
    headers=['Requirement','Specification','Why It Matters'],
    rows=[
        ['Room lighting','Complete darkness — no ambient light','Any ambient light raises the measured black level and reduces apparent contrast, skewing EOTF measurements'],
        ['Screen surface','Clean, no dust or smudges','Contaminants scatter light and affect colorimeter contact readings'],
        ['Room temperature','Stable 20–25°C during full session','Temperature drift affects lamp output and colorimeter readings'],
        ['Screen gain','Note your screen gain value','Gain > 1.0 affects luminance uniformity — measure at screen center'],
        ['Projector warm-up','|B|Minimum 30 minutes|B| with lamp on before any measurement','Lamp color and brightness stabilize after ~20–30 min. Critical for accurate white balance'],
    ],
    col_widths=[1.8,2.0,2.8],hdr_fill=H_GRAY)

heading(doc,'Colorimeter Placement on Projection Screen',2,color=ACCENT,sb=10,sa=4)
body(doc,'The i1Display Pro Plus is a contact colorimeter designed for direct-view displays. For projectors, the following placement method is used:')

mono_block(doc,[
    '  PROJECTION SCREEN  (front view)',
    '',
    '  ┌──────────────────────────────────────────┐',
    '  │                                          │',
    '  │                                          │',
    '  │           ┌───────────────┐              │',
    '  │           │  100% white   │              │',
    '  │           │  test patch   │              │',
    '  │           │               │              │',
    '  │           │    [meter]    │ ← i1Display  │',
    '  │           │    placed     │   Pro Plus   │',
    '  │           │    flat on    │   face-down  │',
    '  │           │    screen     │   on screen  │',
    '  │           └───────────────┘              │',
    '  │                                          │',
    '  └──────────────────────────────────────────┘',
    '                     ▲',
    '              Measure at center of screen only.',
    '              Lens axis should be perpendicular to screen.',
],'F0F0F4')

callout(doc,
    'Place the i1Display Pro Plus flat on the screen surface, sensor face-down against the screen, '
    'with the USB cable running to the side. Use a black cloth or dark foam over the meter body '
    'to block any spill light around the sensor edges. This is standard practice for projector '
    'contact measurement with a colorimeter.',
    fill='EEF4FF',border='006DC6',label='METER PLACEMENT',lcolor=ACCENT)

callout(doc,
    'The i1Display Pro Plus (2000 nit version) has extended filter coverage for wide-color-gamut '
    'displays. For the VW385ES lamp (UHP lamp), use the "LCD (CCFL)" or "Projector" correction '
    'matrix in CalMAN if available, or the generic matrix. The lamp spectrum differs from LED/OLED '
    'displays, which can introduce small metamerism errors — ArgyllPro verification in Phase 5 '
    'will catch any systematic offset.',
    fill='FFFBEA',border='B86C00',label='COLORIMETER NOTE',lcolor=AMBER)

heading(doc,'Network Setup — Connecting All Devices',2,color=ACCENT,sb=10,sa=4)
data_table(doc,
    headers=['Device','Connection','IP Address','Port'],
    rows=[
        ['Raspberry Pi 4 (PGenerator)','Wi-Fi or Ethernet to router','Assign static IP e.g. 192.168.1.50','PGenerator default: 85 (TCP)'],
        ['Calibration PC / Mac (CalMAN)','Wi-Fi or Ethernet to same router','DHCP or static','CalMAN connects out to PGenerator'],
        ['Sony VPL-VW385ES','Ethernet to router','Assign static IP e.g. 192.168.1.100','SDCP: 53484 (for ImageDirector)'],
        ['i1Display Pro Plus','USB to calibration PC','—','Detected by CalMAN and ArgyllPro via USB'],
    ],
    col_widths=[2.0,1.9,1.7,1.0],hdr_fill=H_GRAY)

heading(doc,'Projector Initial Settings  (set before calibration begins)',2,color=ACCENT,sb=10,sa=4)
step_row(doc,1,'Navigate to Picture Menu → Calib. Preset → select Cinema Film 1','This is the preset slot where all HDR calibration will be stored.')
step_row(doc,2,'Set HDR → HDR10  (NOT Auto)','Locks the HDR processing mode. Auto can override Color Space mid-measurement.')
step_row(doc,3,'Set Color Space → BT.2020','Wide gamut container for HDR10 content.')
step_row(doc,4,'Set Gamma Correction → Gamma 8  (will upload curve later)','Placeholder — set to Gamma 8 slot now so it\'s ready after upload.')
step_row(doc,5,'Set Advanced Iris → Full','Maximum dynamic iris range for best contrast.')
step_row(doc,6,'Set Lamp Control → High','HDR requires maximum brightness headroom.')
step_row(doc,7,'Set Motionflow → Off','Disable all processing that could affect measurement stability.')
step_row(doc,8,'Set NR, MPEG NR → Off','Disable noise reduction — these alter the patch values during measurement.')
step_row(doc,9,'Reset Color Correction → all values to 0','Start CMS from a clean baseline.')

# ══════════════════════════════════════════════════════════════
#  PHASE 1 — PQ CURVE UPLOAD
# ══════════════════════════════════════════════════════════════
phase_banner(doc,1,'PQ Curve Generation & Upload','LDT Editor → Sony ImageDirector → Gamma 8 slot')

body(doc,'Before measuring, the PQ tone-mapping curve must be in the projector. This is done first so all subsequent EOTF measurements reflect the actual curve the projector will use.')

heading(doc,'Step 1A — Generate PQ Curve in LDT Editor',2,color=GREEN_DARK,sb=8,sa=4)
step_row(doc,1,'Open LDT Editor in your browser  (npm run dev  or the built app)','Navigate to the Sidebar → HDR PQ section.')
step_row(doc,2,'Set the nit slider to 200  (VW385ES peak brightness)','This is the tone-mapping target. Highlights above 200 nit in the PQ signal will be softly compressed.')
step_row(doc,3,'Click "Generate PQ → 200 nit"','The canvas shows the PQ curve. Verify it bows upward steeply at low end and rolls off toward highlights.')
step_row(doc,4,'Click Export As… → name the file  HDR_PQ_200nit.ldt','Verify the file is exactly 6,656 bytes after saving.')

callout(doc,
    'Optional: Use "Set Ref" then adjust the slider to compare a 150 nit vs 200 nit curve '
    'before committing. The "Preview" ghost curve shows the effect live on the canvas. '
    'For the VW385ES, 200 nit is the recommended starting target.',
    fill='F0FFF4',border='1A6B2E',label='LDT EDITOR TIP',lcolor=GREEN_DARK)

heading(doc,'Step 1B — Upload via Sony ImageDirector',2,color=GREEN_DARK,sb=8,sa=4)
step_row(doc,1,'Open Sony ImageDirector on your PC','Ensure the PC is on the same LAN as the projector.')
step_row(doc,2,'Enter the projector\'s IP address and connect','The VW385ES accepts connections on SDCP protocol. Confirm "Connected" status.')
step_row(doc,3,'Select Upload → Custom Gamma → Gamma 8 slot','Choose the HDR_PQ_200nit.ldt file you exported from LDT Editor.')
step_row(doc,4,'Confirm upload complete — projector display will briefly flash','The Gamma 8 slot now holds the PQ curve.')
step_row(doc,5,'On the projector, confirm Gamma Correction = Gamma 8 is still selected','Navigate to Picture → Expert Setting → Gamma Correction → Gamma 8.')

# ══════════════════════════════════════════════════════════════
#  PHASE 2 — CALMAN + PGENERATOR SETUP
# ══════════════════════════════════════════════════════════════
phase_banner(doc,2,'CalMAN & PGenerator Configuration','Connect CalMAN to meter and pattern generator before measuring')

heading(doc,'CalMAN Workflow Setup',2,color=AMBER,sb=8,sa=4)
data_table(doc,
    headers=['CalMAN Setting','Value','Where to Set'],
    rows=[
        ['Workflow Template','Projector (Manual) or Display Calibration','File → New Workflow → Projector'],
        ['Color Space Target','DCI-P3 D65 / BT.2020 container','Workflow Settings → Color Space'],
        ['White Point Target','D65 (x=0.3127, y=0.3290)','Workflow Settings → White Point'],
        ['EOTF Target','PQ / ST.2084','Workflow Settings → EOTF'],
        ['Peak Luminance','200 nit','Workflow Settings → Max Luminance'],
        ['Black Level','0.002 nit','Workflow Settings → Black Level (measured)'],
        ['dE Formula','dE2000','Workflow Settings → Delta E Formula'],
        ['Meter','X-Rite i1Display Pro Plus','Source → Meter → Select i1DisplayPro'],
        ['Meter Mode','Projector (if available) or LCD CCFL','Source → Meter → Mode'],
        ['Pattern Source','PGenerator','Source → Pattern Generator → PGenerator'],
    ],
    col_widths=[2.0,2.2,2.4],hdr_fill=H_AMBER)

heading(doc,'Connecting CalMAN to PGenerator 1.6',2,color=AMBER,sb=10,sa=4)
step_row(doc,1,'On the Raspberry Pi — launch PGenerator 1.6','PGenerator starts and listens on its default TCP port (85).')
step_row(doc,2,'In CalMAN → Source → Pattern Generator → select PGenerator','Enter the Pi\'s IP address (e.g. 192.168.1.50) and port 85.')
step_row(doc,3,'Click Connect — CalMAN sends a test pattern to verify','You should see a test patch appear on the projection screen.')
step_row(doc,4,'In PGenerator settings, set output to HDR10 / BT.2020 / 10-bit','This ensures patterns are sent as proper HDR10 signal to the projector.')
step_row(doc,5,'Verify projector Information Menu shows correct signal','Menu → Information → confirm: 3840×2160, BT.2020, HDR(HDR10).')

callout(doc,
    'PGenerator outputs HDR10 test patterns at 3840×2160 (4K). The projector will recognize '
    'this as a preset 93–96 or 74–79 signal. Confirm CalMAN is connected and auto-patterns '
    'are advancing correctly before beginning any measurements.',
    fill='FFFBEA',border='B86C00',label='PGENERATOR NOTE',lcolor=AMBER)

# ══════════════════════════════════════════════════════════════
#  PHASE 3 — BASELINE MEASUREMENT
# ══════════════════════════════════════════════════════════════
phase_banner(doc,3,'Baseline Measurement','Measure before any adjustments — documents the starting state')

body(doc,
    'Always measure baseline before touching any projector settings. This documents the '
    'projector\'s out-of-box or previous-calibration state and helps identify which '
    'parameters need the most correction.')

step_row(doc,1,'In CalMAN, run the pre-calibration measurement sequence','Measure: White, Black, R, G, B, Cyan, Magenta, Yellow at 100% stimulus.')
step_row(doc,2,'Run EOTF / Gamma measurement sequence','Measure PQ stimulus levels: 5%, 10%, 20%, 30%, 40%, 50%, 60%, 70%, 80%, 90%, 100%.')
step_row(doc,3,'Record the following baseline values',
    'White luminance (nit) / Black luminance (nit) / White point xy / Primary dE2000 values / EOTF curve shape')
step_row(doc,4,'Save CalMAN session as  "VW385ES_HDR_Baseline.cal"','File → Save. This is your pre-calibration reference.')

callout(doc,
    'Do not close or reset the CalMAN session after baseline. The baseline measurements '
    'are shown alongside post-calibration measurements for direct before/after comparison.',
    fill='EEF4FF',border='006DC6',label='IMPORTANT',lcolor=ACCENT)

# ══════════════════════════════════════════════════════════════
#  PHASE 4 — WHITE BALANCE (COLOR TEMPERATURE)
# ══════════════════════════════════════════════════════════════
phase_banner(doc,4,'White Balance Calibration','Projector: Color Temp → Custom 3  |  CalMAN: White Balance workflow')

body(doc,
    'White balance sets the projector\'s white point to D65 by adjusting the Gain '
    '(white point) and Bias (black point) controls for the three channels. '
    'This must be done before CMS, as the white point affects how all colors are perceived.')

heading(doc,'Two-Point White Balance Method',2,color=PURPLE,sb=8,sa=4)
data_table(doc,
    headers=['Adjustment','Controls','Pattern Used','Target'],
    rows=[
        ['High luminance (Gain)','Gain R, Gain G, Gain B  in Color Temp → Custom 3','100% white patch','D65 white point: x=0.3127, y=0.3290'],
        ['Low luminance (Bias)','Bias R, Bias G, Bias B  in Color Temp → Custom 3','20% white patch (near black)','D65 — same xy target, minimum dE'],
    ],
    col_widths=[1.8,2.2,1.7,1.9],hdr_fill=H_PURPLE)

step_row(doc,1,'In CalMAN, navigate to White Balance workflow','Select the two-point white balance page.')
step_row(doc,2,'Project 100% white patch via PGenerator','CalMAN auto-advances to this pattern.')
step_row(doc,3,'Adjust Gain R / G / B on projector until dE < 2.0','Navigate: Picture → Color Temp → Custom 3. Adjust Gain R, G, B while watching CalMAN dE reading live. |B|Gain G is the reference — adjust R and B relative to G.|B|')
step_row(doc,4,'Project 20% white patch via PGenerator','Switch to low-luminance pattern.')
step_row(doc,5,'Adjust Bias R / G / B until dE < 2.0','Small adjustments — Bias has less range than Gain. Aim for neutral gray at low luminance.')
step_row(doc,6,'Iterate Gain → Bias → Gain until both are stable','Gain and Bias interact. Typically 2–3 iterations converges.')
step_row(doc,7,'Record final Gain and Bias values for all three channels','Document in CalMAN notes or a spreadsheet.')

callout(doc,
    'White point dE2000 < 2.0 is acceptable, < 1.0 is excellent. '
    'If you cannot reach < 2.0 on Bias without pushing Gain out of range, prioritize '
    'the Gain (100% white) setting. Low-luminance white balance has less visual impact '
    'on HDR content than the high-luminance white point.',
    fill='EEF4FF',border='5B2D8E',label='WHITE BALANCE TOLERANCE',lcolor=PURPLE)

# ══════════════════════════════════════════════════════════════
#  PHASE 5 — EOTF VERIFICATION
# ══════════════════════════════════════════════════════════════
phase_banner(doc,5,'EOTF / PQ Curve Verification','Confirm the uploaded PQ curve tracks the ST.2084 target')

body(doc,
    'The PQ curve uploaded to Gamma 8 should closely follow the ST.2084 (PQ) EOTF shape '
    'from 0 to 200 nit. CalMAN measures the luminance output at multiple stimulus levels '
    'and compares against the mathematical PQ target.')

step_row(doc,1,'In CalMAN, navigate to EOTF / Gamma measurement page','Select PQ as the target EOTF.')
step_row(doc,2,'Run the automated EOTF sequence via PGenerator','CalMAN sends stimulus patches at: 0%, 5%, 10%, 20%, 30%, 40%, 50%, 60%, 70%, 80%, 90%, 95%, 100%.')
step_row(doc,3,'Review the EOTF tracking curve in CalMAN','The measured curve should track the PQ target closely. Deviations in the 50–80% range indicate the tone-mapping target nit level may need adjustment.')
step_row(doc,4,'If average deviation > 8%, regenerate curve at adjusted nit target',
    'If measured luminance at 100% stimulus is lower than 200 nit (e.g. 180 nit), '
    'regenerate the PQ curve in LDT Editor with the measured actual peak (180 nit), '
    're-export, and re-upload via ImageDirector.')
step_row(doc,5,'Re-run EOTF measurement to confirm improvement','Repeat until average deviation < 5%.')

mono_block(doc,[
    '  PQ EOTF SHAPE — What CalMAN should show:',
    '',
    '  Output (nit)',
    '  200 ┤                                          ●  ← 100% stimulus = ~200 nit',
    '      │                                      ●',
    '      │                                  ●',
    '  100 ┤                             ●',
    '      │                        ●',
    '      │                   ●',
    '   50 ┤              ●',
    '      │         ●',
    '      │    ●',
    '    0 ┤●',
    '      └──────────────────────────────────────────',
    '        0%   20%   40%   60%   80%  100%   Stimulus',
    '',
    '  The curve should accelerate in the low range and flatten near peak.',
    '  Deviations > 8% at any point warrant a curve re-upload.',
],'1E1E2E',title='Expected EOTF Shape — CalMAN Graph Reference')

# ══════════════════════════════════════════════════════════════
#  PHASE 6 — CMS GAMUT CALIBRATION
# ══════════════════════════════════════════════════════════════
phase_banner(doc,4,'CMS Gamut Calibration','CalMAN measures primaries/secondaries  →  Manual entry into projector Color Correction')

body(doc,
    'The Color Management System (CMS) corrects the chromaticity of each primary and '
    'secondary color to hit P3-D65 targets. CalMAN measures each color, shows the error, '
    'and you manually adjust Hue / Saturation / Brightness in the projector menu. '
    'This is an iterative process — typically 2–4 rounds.')

heading(doc,'Calibration Order — Primaries Before Secondaries',2,color=PURPLE,sb=8,sa=4)
data_table(doc,
    headers=['Round','Colors to Adjust','Reason'],
    rows=[
        ['Round 1','Red, Green, Blue  (primaries first)','Primaries define the gamut triangle corners. Get these close before touching secondaries.'],
        ['Round 2','Cyan, Magenta, Yellow  (secondaries)','Secondaries are blends of two primaries — calibrate primaries first or secondary corrections will shift.'],
        ['Round 3','Re-check all 6 colors','Primary corrections can slightly affect secondaries and vice versa. A verification round catches cross-influence.'],
        ['Round 4 (if needed)','Fine-tune any outlier > dE 3.0','Focus only on colors that are still out of tolerance.'],
    ],
    col_widths=[0.9,2.4,3.3],hdr_fill=H_PURPLE)

heading(doc,'Per-Color Adjustment Workflow',2,color=PURPLE,sb=10,sa=4)
step_row(doc,1,'In CalMAN, project 100% Red patch via PGenerator','CalMAN shows measured vs target for Red: dE2000, Hue error (°), Saturation error (%), Luminance error (%).')
step_row(doc,2,'On the projector: Picture → Expert Setting → Color Correction → On → Color Select → Red','Navigate to the Red color axis. You will see Hue, Saturation, Brightness sub-controls.')
step_row(doc,3,'Adjust Hue first — watch CalMAN dE2000 update in real time','Hue error corrects angular position on the color wheel. Positive = shift toward yellow; Negative = shift toward magenta.')
step_row(doc,4,'Adjust Saturation — correct the distance from white point','Positive = more saturated (more vivid); Negative = less saturated (closer to white).')
step_row(doc,5,'Adjust Brightness — correct luminance','Only adjust if CalMAN shows a significant luminance error (> 5%). Brightness changes can affect the overall gamma response.')
step_row(doc,6,'Repeat for Green, Blue, then Cyan, Magenta, Yellow','Follow the same Hue → Saturation → Brightness sequence for each color.')
step_row(doc,7,'After all 6 colors, run a full re-measurement','CalMAN measures all primaries and secondaries in sequence. Check for any new outliers introduced by cross-influence.')

callout(doc,
    'The projector CMS only adjusts colors at 100% saturation. Lower saturation levels '
    '(25%, 50%, 75%) are affected by the correction but not perfectly — expect some '
    'residual error at mid-saturations. This is a known limitation of a 6-axis CMS '
    'and is acceptable for home cinema use (dE < 5 at 50% saturation is typical).',
    fill='FFF0F0',border='991A1A',label='CMS LIMITATION',lcolor=RED_DARK)

heading(doc,'CMS Quick Reference — Which Control Fixes Which Error',2,color=PURPLE,sb=10,sa=4)
data_table(doc,
    headers=['CalMAN Shows This Error','Adjust This Control','Direction'],
    rows=[
        ['Hue angle off target (° error)','Hue','Move in the direction CalMAN indicates — watch the xy chromaticity point move'],
        ['Saturation too high (color too vivid)','Saturation','Decrease (negative direction)'],
        ['Saturation too low (color washed out)','Saturation','Increase (positive direction)'],
        ['Luminance too high','Brightness','Decrease'],
        ['Luminance too low','Brightness','Increase'],
        ['dE still high after all 3 adjustments','Re-check white balance first','White balance errors project into all color measurements — re-verify Gain/Bias if CMS won\'t converge'],
    ],
    col_widths=[2.4,1.7,2.5],hdr_fill=H_PURPLE)

# ══════════════════════════════════════════════════════════════
#  PHASE 7 — ARGYLLPRO VERIFICATION
# ══════════════════════════════════════════════════════════════
phase_banner(doc,5,'ArgyllPro ColorMeter v2 — Independent Verification','Cross-check CalMAN results with a separate software stack')

body(doc,
    'ArgyllPro ColorMeter provides an independent verification pass. Because it uses '
    'different internal algorithms than CalMAN, agreement between the two confirms '
    'the calibration is solid and not an artifact of a single software\'s measurement model.')

heading(doc,'Why Use ArgyllPro as a Second Tool',2,color=TEAL,sb=8,sa=4)
data_table(doc,
    headers=['ArgyllPro Capability','How It Helps'],
    rows=[
        ['Independent meter driver (ArgyllCMS engine)','Reads the same i1Display Pro Plus through a different software stack — any systematic CalMAN error shows as a discrepancy'],
        ['Saturation sweep at 25/50/75/100%','Measures each primary and secondary at multiple saturation levels — reveals how well the 100% CMS correction translates to lower saturations'],
        ['Detailed dE report export','Generates a full per-patch dE2000 report for documentation and client records'],
        ['Display ICC profile creation','Profiles the projector\'s post-calibration state — useful for understanding the display\'s gamut coverage in a standard format'],
        ['Spot measurement mode','Quick individual patch reads without a full sequence — useful for spot-checking specific colors during CMS work'],
    ],
    col_widths=[2.2,4.4],hdr_fill=H_TEAL)

heading(doc,'ArgyllPro Verification Workflow',2,color=TEAL,sb=10,sa=4)
step_row(doc,1,'Connect i1Display Pro Plus to PC via USB','ArgyllPro auto-detects it on launch. Confirm meter name shown in ArgyllPro device list.')
step_row(doc,2,'In ArgyllPro → New Session → Display Measurement','Select "Projector" as display type if available, otherwise "LCD".')
step_row(doc,3,'Set color space to Rec.2020 / BT.2020, white point D65, EOTF PQ',
    'Match these to the CalMAN session targets so results are directly comparable.')
step_row(doc,4,'Run spot measurements on primary and secondary patches',
    'Use PGenerator to hold each patch steady. In ArgyllPro, click Measure for each patch. '
    'Measure: White, Black, R, G, B, C, M, Y at 100% stimulus.')
step_row(doc,5,'Run saturation sweep — measure each primary at 25%, 50%, 75%, 100%',
    'Use PGenerator to send each saturation level. Measure Red at 25%→50%→75%→100%, '
    'then repeat for Green, Blue. This reveals how the CMS correction translates downward.')
step_row(doc,6,'Compare ArgyllPro dE2000 results against CalMAN results',
    'Values should agree within ± 0.5 dE. Larger discrepancies suggest a meter placement '
    'or correction matrix issue — re-seat the colorimeter and re-measure.')
step_row(doc,7,'Export ArgyllPro report as PDF or CSV','File → Export Report. Save as VW385ES_HDR_ArgyllVerification.')

callout(doc,
    'If ArgyllPro and CalMAN show consistently different dE values for the same patches '
    '(e.g. ArgyllPro reads 1.5 dE higher across all colors), this suggests a correction '
    'matrix mismatch between the two software tools for this specific lamp type. '
    'In this case, trust the relative improvements shown by each tool separately, '
    'and document both readings.',
    fill='D0FFF8',border='0E7A72',label='DISCREPANCY NOTE',lcolor=TEAL)

heading(doc,'ArgyllCMS Command Reference  (advanced / command-line users)',2,color=TEAL,sb=10,sa=4)
body(doc,'ArgyllPro ColorMeter v2 provides a GUI for these operations, but the underlying ArgyllCMS commands can be run directly for more control:',sa=4)

mono_block(doc,[
    '# Single spot measurement (hold patch steady in PGenerator):',
    'spotread -d 1 -c 1 -e',
    '',
    '# Full display read session (reads a set of patches):',
    'dispread -d 1 -c 1 -k calibration.cal display_test',
    '',
    '# Create ICC profile from measurements:',
    'colprof -v -D "VW385ES_HDR" -q m -a S display_test',
    '',
    '# Verify ICC profile against measurements:',
    'profcheck -v display_test.ti3 VW385ES_HDR.icc',
    '',
    '# Notes:',
    '#   -d 1  = display number (adjust if multiple monitors)',
    '#   -c 1  = communication port / meter index',
    '#   -e    = extra emit measure (for emissive display)',
    '#   Place meter on screen before running spotread',
],'1E1E2E',title='ArgyllCMS Command-Line Reference')

# ══════════════════════════════════════════════════════════════
#  TOLERANCE GUIDE
# ══════════════════════════════════════════════════════════════
heading(doc,'Tolerance Reference Guide',1,sb=16)
body(doc,'Use this table to evaluate whether each measurement result is acceptable. Re-calibrate any metric in the "Action Required" column.',sa=6)
tolerance_table(doc)

callout(doc,
    'Saturation sweep dE values at 50% saturation will be higher than 100% saturation — '
    'this is normal for a 6-axis CMS. A result of dE < 5 at 50% saturation is acceptable '
    'for home cinema. dE > 8 at 50% saturation warrants a second CMS pass focusing on '
    'the specific offending color axis.',
    fill='D0FFF8',border='0E7A72',label='SATURATION SWEEP NOTE',lcolor=TEAL)

# ══════════════════════════════════════════════════════════════
#  PHASE 8 — SAVE & DOCUMENT
# ══════════════════════════════════════════════════════════════
phase_banner(doc,6,'Save, Document & Verify Recall','Final steps — store all settings and confirm recall works correctly')

step_row(doc,1,'Confirm projector Calib. Preset is set to Cinema Film 1','All settings adjusted during calibration are automatically saved to this preset.')
step_row(doc,2,'Power-cycle the projector and let it warm up for 10 minutes','Confirms that settings survived the power cycle and recall correctly.')
step_row(doc,3,'In CalMAN, run a final post-calibration measurement pass','Measure White, Black, R, G, B, C, M, Y and EOTF. This is your final "after" record.')
step_row(doc,4,'Save CalMAN session as  VW385ES_HDR_Calibrated_[date].cal','This file preserves before/after measurements, dE values, and all settings.')
step_row(doc,5,'Save ArgyllPro session and export final report','File → Export. Include both saturation sweep and primary/secondary dE table.')
step_row(doc,6,'Document all manual settings in a reference sheet',
    'Record: Gain R/G/B, Bias R/G/B, CMS Hue/Sat/Brightness per color, '
    'Gamma slot (8 = PQ 200 nit), Color Space (BT.2020), Lamp level, Iris mode.')
step_row(doc,7,'Confirm Calib. Preset switching works correctly',
    'Switch to Reference (SDR 709) then back to Cinema Film 1 (HDR). '
    'Verify all settings recall correctly including Gamma Correction slot selection.')

# ─── CALMAN SESSION CHECKLIST ───
heading(doc,'CalMAN Session Checklist',2,color=ACCENT,sb=12,sa=4)
data_table(doc,
    headers=['CalMAN File to Save','Contains','When to Save'],
    rows=[
        ['VW385ES_HDR_Baseline.cal','Pre-calibration measurements — all primaries, EOTF','End of Phase 3'],
        ['VW385ES_HDR_WhiteBalance.cal','Post white balance measurements','End of Phase 4'],
        ['VW385ES_HDR_Calibrated_[date].cal','Full before/after comparison, final dE values','End of Phase 8'],
    ],
    col_widths=[2.6,2.6,1.4],hdr_fill=H_ACCENT)

# ─── FULL WORKFLOW SUMMARY ───
heading(doc,'Complete Calibration Workflow Summary',1,sb=14)
mono_block(doc,[
    '  BEFORE STARTING',
    '  ├── Dark room, 30-min projector warm-up, meter on screen',
    '  ├── PGenerator on RPi4 → HDMI → VW385ES',
    '  ├── CalMAN connected to PGenerator via LAN',
    '  └── Projector: Cinema Film 1, HDR10, BT.2020, Gamma 8, Iris Full',
    '',
    '  PHASE 1 — PQ CURVE  (LDT Editor + ImageDirector)',
    '  └── Generate PQ 200 nit → export .ldt → upload to Gamma 8',
    '',
    '  PHASE 2 — CALMAN SETUP',
    '  └── Configure meter, PGenerator, targets (P3-D65, PQ, 200 nit)',
    '',
    '  PHASE 3 — BASELINE  (CalMAN)',
    '  └── Measure all colors + EOTF before touching anything → Save .cal',
    '',
    '  PHASE 4 — WHITE BALANCE  (CalMAN + projector Custom 3)',
    '  └── 100% white → Gain R/G/B → 20% white → Bias R/G/B → iterate',
    '',
    '  PHASE 5 — EOTF CHECK  (CalMAN)',
    '  └── Measure PQ tracking → if deviation > 8%, re-upload curve',
    '',
    '  PHASE 6 — CMS  (CalMAN + projector Color Correction)',
    '  ├── Round 1: Red → Green → Blue',
    '  ├── Round 2: Cyan → Magenta → Yellow',
    '  └── Round 3: Full re-check, fix outliers',
    '',
    '  PHASE 7 — ARGYLLPRO VERIFICATION',
    '  ├── Independent primary/secondary measurement',
    '  ├── Saturation sweep: 25/50/75/100% per primary',
    '  └── Export dE report',
    '',
    '  PHASE 8 — SAVE & CONFIRM',
    '  └── Power-cycle → recall test → final CalMAN pass → document all',
],'F0F4F8')

footer(doc,'Sony VPL-VW385ES  ·  HDR Calibration Procedure  ·  CalMAN Business 5.12 + i1Display Pro Plus + PGenerator 1.6 + ArgyllPro v2')

out='/Users/johnlee/Downloads/sony-ldt-editor/Resources/Report3_HDR_Calibration_Procedure.docx'
doc.save(out)
print(f'Saved: {out}')
