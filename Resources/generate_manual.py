from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Page margins ───
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(0.85)
section.bottom_margin = Inches(0.85)

# ─── Palette ───
BLACK      = RGBColor(0x1A, 0x1A, 0x1A)
DARK_GRAY  = RGBColor(0x44, 0x44, 0x44)
MID_GRAY   = RGBColor(0x77, 0x77, 0x77)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
RULE_GRAY  = RGBColor(0xCC, 0xCC, 0xCC)
SONY_BLUE  = RGBColor(0x00, 0x3B, 0x8E)   # Sony brand blue
ACCENT     = RGBColor(0x00, 0x6D, 0xC6)   # lighter accent
HDR_AMBER  = RGBColor(0xB8, 0x6C, 0x00)
HDR_BG     = RGBColor(0xFF, 0xF8, 0xEC)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
ROW_ALT    = RGBColor(0xF7, 0xF9, 0xFF)   # very faint blue tint for alt rows
ROW_HEAD   = RGBColor(0x00, 0x3B, 0x8E)   # table header = sony blue

# ─── Helper: set paragraph shading ───
def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

# ─── Helper: shade a table cell ───
def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ─── Helper: set cell border ───
def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'),   kwargs[edge].get('val','single'))
            tag.set(qn('w:sz'),    kwargs[edge].get('sz','4'))
            tag.set(qn('w:space'),'0')
            tag.set(qn('w:color'), kwargs[edge].get('color','auto'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

# ─── Helper: remove borders from a table ───
def remove_table_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'none')
        tblBorders.append(tag)
    tblPr.append(tblBorders)

# ─── Helper: add a horizontal rule paragraph ───
def add_rule(color_hex='CCCCCC', thickness=4):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(thickness))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    return p

# ─── Helper: styled heading ───
def add_heading(text, level=1, color=SONY_BLUE, size=None, space_before=18, space_after=6, bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    else:
        sizes = {1: 20, 2: 14, 3: 12, 4: 11}
        run.font.size = Pt(sizes.get(level, 11))
    run.font.name = 'Calibri'
    return p

# ─── Helper: body paragraph ───
def add_body(text, color=DARK_GRAY, size=10, italic=False, space_before=2, space_after=4):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name   = 'Calibri'
    run.font.size   = Pt(size)
    run.font.color.rgb = color
    run.italic      = italic
    return p

# ─── Helper: styled two-column table ───
def add_table(headers, rows, col_widths=None, alt_rows=True, note=None):
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(table)

    # header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        shade_cell(cell, '003B8E')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # thin white bottom border to separate header
        set_cell_border(cell, bottom={'val':'single','sz':'4','color':'FFFFFF'})
        p   = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(h.upper())
        run.bold = True
        run.font.name  = 'Calibri'
        run.font.size  = Pt(9)
        run.font.color.rgb = WHITE

    # data rows
    for ri, row_data in enumerate(rows):
        tr   = table.rows[ri + 1]
        fill = 'F7F9FF' if (alt_rows and ri % 2 == 1) else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = tr.cells[ci]
            shade_cell(cell, fill)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell, bottom={'val':'single','sz':'2','color':'E0E0E0'})
            p   = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            # first column bold
            parts = cell_text.split('||bold||') if '||bold||' in cell_text else [cell_text]
            for pi, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name  = 'Calibri'
                run.font.size  = Pt(9.5)
                run.font.color.rgb = BLACK if ci == 0 else DARK_GRAY
                if ci == 0:
                    run.bold = True

    # column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    if note:
        pn = doc.add_paragraph()
        pn.paragraph_format.space_before = Pt(3)
        pn.paragraph_format.space_after  = Pt(8)
        rn = pn.add_run('Note: ' + note)
        rn.font.name   = 'Calibri'
        rn.font.size   = Pt(8.5)
        rn.font.color.rgb = MID_GRAY
        rn.italic = True

    return table

# ─── Helper: callout box ───
def add_callout(text, fill_hex='FFF8EC', border_hex='B86C00', label=None, label_color=HDR_AMBER):
    p    = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    shade_paragraph(p, fill_hex)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for edge in ('top','left','bottom','right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   'single')
        tag.set(qn('w:sz'),    '12' if edge == 'left' else '4')
        tag.set(qn('w:space'), '4')
        tag.set(qn('w:color'), border_hex)
        pBdr.append(tag)
    pPr.append(pBdr)
    if label:
        rl = p.add_run(label + '  ')
        rl.bold = True
        rl.font.name  = 'Calibri'
        rl.font.size  = Pt(9.5)
        rl.font.color.rgb = label_color
    rb = p.add_run(text)
    rb.font.name  = 'Calibri'
    rb.font.size  = Pt(9.5)
    rb.font.color.rgb = DARK_GRAY
    return p

# ═══════════════════════════════════════════════════════════
#  COVER BLOCK
# ═══════════════════════════════════════════════════════════
p_cover = doc.add_paragraph()
p_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cover.paragraph_format.space_before = Pt(6)
p_cover.paragraph_format.space_after  = Pt(2)
shade_paragraph(p_cover, '003B8E')
rc = p_cover.add_run('VPL-VW385ES  ·  Picture & Color Calibration Controls')
rc.font.name  = 'Calibri'
rc.font.size  = Pt(17)
rc.font.color.rgb = WHITE
rc.bold = True

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after  = Pt(4)
shade_paragraph(p_sub, '003B8E')
rs = p_sub.add_run('Sony 4K SXRD Home Cinema Projector  ·  Verified from Operating Instructions (47254410M)')
rs.font.name  = 'Calibri'
rs.font.size  = Pt(10)
rs.font.color.rgb = RGBColor(0xAA, 0xC8, 0xFF)

add_rule('1A6FD8', 8)

# ═══════════════════════════════════════════════════════════
#  SECTION 1 — CALIBRATED PRESET (Picture Modes)
# ═══════════════════════════════════════════════════════════
add_heading('1  Calibrated Preset  (Picture Modes)', level=1)
add_body(
    'Nine picture modes selectable via the CALIBRATED PRESET buttons on the remote or the Calib. Preset '
    'item in the Picture Menu. Each mode stores independent settings per input connector.',
    space_after=8
)

add_table(
    headers=['Mode', 'Intended Use'],
    rows=[
        ['Cinema Film 1', 'Reproduces the dynamic, clear look of master positive film — primary cinema mode'],
        ['Cinema Film 2', 'Rich tone and color of a movie theater; based on Cinema Film 1'],
        ['Reference',     'Faithful reproduction of source material; no processing applied — ideal for calibration'],
        ['TV',            'TV programs, sports, concerts, general video content'],
        ['Photo',         'Still images from a digital camera'],
        ['Game',          'Optimised for gaming; well-modulated colours and fast response'],
        ['Bright Cinema', 'Movie watching in a bright room such as a living room'],
        ['Bright TV',     'TV/sports in a bright room'],
        ['User',          'Fully user-customisable; factory default matches Reference'],
    ],
    col_widths=[1.8, 4.6],
)

add_callout(
    'All picture quality adjustments (Contrast, Brightness, Color, Gamma, etc.) are saved independently '
    'for each input connector and each Calib. Preset mode.',
    fill_hex='EEF4FF', border_hex='006DC6', label='TIP', label_color=ACCENT
)

# ═══════════════════════════════════════════════════════════
#  SECTION 2 — PICTURE MENU
# ═══════════════════════════════════════════════════════════
add_heading('2  Picture Menu', level=1, space_before=14)
add_body('Primary adjustment controls accessible from the main Picture menu.', space_after=8)

add_heading('2.1  Image Processing', level=2, color=ACCENT, space_before=10, space_after=4)
add_table(
    headers=['Control', 'Options', 'Notes'],
    rows=[
        ['Reality Creation',   'On / Off',                       'Super-resolution. Sub-controls: Database (Normal / Mastered in 4K), Resolution level, Noise Filtering level'],
        ['Advanced Iris',      'Full / Limited / Off',           'Dynamic iris under Cinema Black Pro. Full = maximum dynamic range; Limited = slower, darker, better for very dark rooms'],
        ['Contrast Enhancer',  'High / Middle / Low / Off',      'Automatically corrects bright/dark balance per scene to optimise contrast and perceived sharpness'],
        ['Lamp Control',       'High / Low',                     'High = maximum brightness; Low = deeper blacks, quieter fan, longer lamp life'],
        ['Motionflow',         'Impulse / Combination / Smooth High / Smooth Low / True Cinema / Off', 'Frame interpolation engine. True Cinema preserves 24fps cadence exactly'],
    ],
    col_widths=[1.6, 2.1, 2.7],
)

add_heading('2.2  Core Tone & Color Controls', level=2, color=ACCENT, space_before=10, space_after=4)
add_table(
    headers=['Control', 'Range / Options', 'Notes'],
    rows=[
        ['Contrast',   'Slider  (HDR: "Contrast(HDR)")', 'White-point adjustment. When HDR10/HLG/Auto HDR is active, label changes to Contrast(HDR)'],
        ['Brightness', 'Slider',                          'Black-lift adjustment. Set this first, then configure Dynamic Control for room brightness'],
        ['Color',      'Slider',                          'Master saturation — affects all hues equally'],
        ['Hue',        'Slider (–red / +green)',          'Master hue rotation'],
        ['Sharpness',  'Slider',                          'Remote shortcut: SHARPNESS +/–. Higher = sharper edges; Lower = softer, less noise'],
    ],
    col_widths=[1.6, 2.1, 2.7],
)

add_heading('2.3  Color Temperature', level=2, color=ACCENT, space_before=10, space_after=4)
add_body(
    'Four fixed presets plus five fully user-adjustable custom slots. '
    'Each slot stores independent Gain and Bias values for all three channels.',
    space_after=6
)
add_table(
    headers=['Preset', 'Colour Temp.', 'Character'],
    rows=[
        ['D93', '~9 300 K', 'Blue-tinted whites; standard TV colour temperature in Japan'],
        ['D75', '~7 500 K', 'Neutral tint between D93 and D65'],
        ['D65', '~6 500 K', 'Standard illuminant for broadcast and cinema mastering'],
        ['D55', '~5 500 K', 'Warm, reddish whites'],
        ['Custom 1', 'Factory = D93', 'User-adjustable'],
        ['Custom 2', 'Factory = D75', 'User-adjustable'],
        ['Custom 3', 'Factory = D65', 'User-adjustable'],
        ['Custom 4', 'Factory = D55', 'User-adjustable'],
        ['Custom 5', 'Brightness priority', 'User-adjustable'],
    ],
    col_widths=[1.3, 1.5, 3.6],
)

add_body('Each of the 9 colour temperature entries (D93–D55 and Custom 1–5) exposes 6 sub-controls:', space_before=6, space_after=4)
add_table(
    headers=['Sub-Control', 'Affects', 'Function'],
    rows=[
        ['Gain R', 'White point — Red',   'Sets how red the peak-white appears'],
        ['Gain G', 'White point — Green', 'Sets how green the peak-white appears'],
        ['Gain B', 'White point — Blue',  'Sets how blue the peak-white appears'],
        ['Bias R', 'Black point — Red',   'Lifts or lowers the red channel at black'],
        ['Bias G', 'Black point — Green', 'Lifts or lowers the green channel at black'],
        ['Bias B', 'Black point — Blue',  'Lifts or lowers the blue channel at black'],
    ],
    col_widths=[1.3, 1.8, 3.3],
)
add_callout(
    '6 sub-controls × 9 colour temperature slots = up to 54 individual colour temperature adjustments.',
    fill_hex='EEF4FF', border_hex='006DC6', label='NOTE', label_color=ACCENT
)

# ═══════════════════════════════════════════════════════════
#  SECTION 3 — EXPERT SETTING (sub-menu)
# ═══════════════════════════════════════════════════════════
add_heading('3  Expert Setting  (Picture Menu Sub-Section)', level=1, space_before=14)
add_body(
    'Accessed within the Picture Menu. These controls are hidden at the basic level and require '
    'navigating into the Expert Setting sub-menu.',
    space_after=8
)

add_heading('3.1  Noise & Processing', level=2, color=ACCENT, space_before=10, space_after=4)
add_table(
    headers=['Control', 'Options', 'Notes'],
    rows=[
        ['NR (Noise Reduction)',       'Auto / High / Middle / Low / Off', 'Reduces analogue-style noise and picture roughness'],
        ['MPEG NR',                    'Auto / High / Middle / Low / Off', 'Targets block noise and mosquito noise from digital compression'],
        ['Smooth Gradation',           'High / Middle / Low / Off',        'Smooths banding in flat gradient areas of the image'],
        ['Film Mode',                  'Auto / Off',                       'Detects and handles film-sourced progressive content'],
    ],
    col_widths=[1.8, 2.0, 2.6],
)

add_heading('3.2  Gamma Correction', level=2, color=ACCENT, space_before=10, space_after=4)
add_body(
    'Eleven selectable gamma settings plus Off. The six standard power-law presets cover '
    'the broadcast/cinema range; four custom slots (Gamma 7–10) are uploaded via Sony '
    'ImageDirector software as .ldt files.',
    space_after=6
)
add_table(
    headers=['Setting', 'Type', 'Character'],
    rows=[
        ['1.8',      'Power law', 'Brightest overall — suitable for very bright rooms'],
        ['2.0',      'Power law', 'Moderately bright'],
        ['2.1',      'Power law', 'Between 2.0 and 2.2'],
        ['2.2',      'Power law', 'Broadcast SDR standard (sRGB / BT.709 display gamma)'],
        ['2.4',      'Power law', 'BT.1886 ideal — darker mid-tones, suited for dark rooms'],
        ['2.6',      'Power law', 'Darkest overall — maximum shadow detail visibility'],
        ['Gamma 7',  'Custom LDT', 'Simulates film gamma curve (Sony description). Upload via ImageDirector'],
        ['Gamma 8',  'Custom LDT', 'Increases image sharpness; for bright viewing environments'],
        ['Gamma 9',  'Custom LDT', 'Brighter than Gamma 8; for bright viewing environments'],
        ['Gamma 10', 'Custom LDT', 'Increases sharpness; for watching TV in bright rooms'],
        ['Off',      '—',          'No gamma correction applied'],
    ],
    col_widths=[1.2, 1.3, 3.9],
)
add_callout(
    'Gamma 7–10 are the custom slots targeted by this LDT editor. '
    'The Sony descriptions above are factory defaults; the slots accept any uploaded curve.',
    fill_hex='FFF8EC', border_hex='B86C00', label='EDITOR NOTE', label_color=HDR_AMBER
)

add_heading('3.3  Color Correction', level=2, color=ACCENT, space_before=10, space_after=4)
add_body(
    'A per-axis 6-colour correction matrix. For each colour axis, three properties can be adjusted '
    'independently while watching the projected image.',
    space_after=6
)
add_table(
    headers=['Colour Axis', 'Hue', 'Saturation', 'Brightness'],
    rows=[
        ['Red',     '✓', '✓', '✓'],
        ['Yellow',  '✓', '✓', '✓'],
        ['Green',   '✓', '✓', '✓'],
        ['Cyan',    '✓', '✓', '✓'],
        ['Blue',    '✓', '✓', '✓'],
        ['Magenta', '✓', '✓', '✓'],
    ],
    col_widths=[1.6, 1.5, 1.5, 1.8],
)
add_callout(
    '6 colour axes × 3 properties = 18 individual colour correction adjustments.',
    fill_hex='EEF4FF', border_hex='006DC6', label='NOTE', label_color=ACCENT
)

add_heading('3.4  Color Space', level=2, color=ACCENT, space_before=10, space_after=4)
add_body(
    'Selects or customises the colour gamut. When "Custom" is selected, a per-colour '
    'chromaticity adjustment sub-menu becomes available.',
    space_after=6
)
add_table(
    headers=['Mode', 'Description'],
    rows=[
        ['BT.709',        'ITU-R BT.709 — standard for HD broadcast and Blu-ray. Equivalent to sRGB'],
        ['BT.2020',       'Wide colour gamut for HDR content. Automatically selected when HDR Auto detects a BT.2020 signal'],
        ['Color Space 1', 'Sony-tuned preset for TV programs, sports, and concert video'],
        ['Color Space 2', 'Sony-tuned preset for TV/sports in a bright room'],
        ['Color Space 3', 'Sony-tuned preset for movies in a bright room'],
        ['Custom',        'Full user control — exposes the chromaticity adjustment sub-menu below'],
    ],
    col_widths=[1.6, 4.8],
)

add_heading('Color Space Gamut Adjustment', level=3,
            color=DARK_GRAY, size=10, space_before=8, space_after=4, bold=True)
add_body(
    'When any Color Space mode is active, two global sliders appear that shift the '
    'entire selected gamut as a whole. These are not per-primary adjustments — they '
    'apply equally across the full colour space.',
    space_after=6
)
add_table(
    headers=['Slider', 'Direction', 'Effect'],
    rows=[
        ['Cyan – Red',    'Negative → Cyan  |  Positive → Red',   'Shifts the entire gamut along the cyan–red opponent axis'],
        ['Magenta – Green','Negative → Magenta  |  Positive → Green','Shifts the entire gamut along the magenta–green opponent axis'],
    ],
    col_widths=[1.6, 2.4, 2.4],
    note='These 2 sliders are present for all Color Space modes (BT.709, BT.2020, '
         'Color Space 1–3, Custom). They function as a global gamut white-point trim, '
         'not individual primary/secondary colour adjustments.'
)
add_callout(
    '2 sliders total for the entire color space — confirmed from OSD screenshot. '
    'Earlier documentation suggesting per-primary adjustment (6 × 2 = 12) was incorrect.',
    fill_hex='FFF8EC', border_hex='B86C00', label='VERIFIED OSD', label_color=HDR_AMBER
)

add_heading('3.5  Remaining Expert Controls', level=2, color=ACCENT, space_before=10, space_after=4)
add_table(
    headers=['Control', 'Options', 'Notes'],
    rows=[
        ['Clear White',         'High / Low / Off',       'Emphasises vivid whites'],
        ['x.v.Color',           'On / Off',               'For x.v.Color-capable sources. Enabling this disables gamma correction'],
        ['HDR',                 'Auto / HDR10 / HLG / Off', 'Auto detects HDR metadata and sets optimal processing. BT.2020 Color Space is set automatically when Auto detects a BT.2020 signal'],
        ['Input Lag Reduction', 'On / Off',               'Reduces display latency. When On, Motionflow, NR, and MPEG NR are disabled'],
    ],
    col_widths=[1.8, 2.1, 2.5],
)

# ═══════════════════════════════════════════════════════════
#  SECTION 4 — ADVANCED PICTURE MENU
# ═══════════════════════════════════════════════════════════
add_heading('4  Advanced Picture Menu', level=1, space_before=14)
add_body(
    'Separate menu section for panel-level auto calibration. Compensates for colour drift '
    'that occurs after extended use of the SXRD panels.',
    space_after=8
)
add_table(
    headers=['Function', 'Description'],
    rows=[
        ['Pre Check', 'Measures current colour against factory reference. Reports dE (colour difference) before correction'],
        ['Adjust',    'Performs automatic calibration. Lens returns to factory default position during process (~few minutes)'],
        ['Before/After', 'Toggles between factory and post-calibration settings at a fixed frequency so you can evaluate the effect on a live image'],
        ['Reset',     'Discards calibration results and reverts to factory default panel settings'],
    ],
    col_widths=[1.6, 4.8],
    note='Perform calibration after 30+ minutes warm-up. Avoid turning off power or operating controls during Pre Check/Adjust.'
)

# ═══════════════════════════════════════════════════════════
#  SECTION 5 — SUMMARY COUNT
# ═══════════════════════════════════════════════════════════
add_heading('5  Total Adjustment Count Summary', level=1, space_before=14)
add_body('Approximate total of distinct calibration-relevant adjustments, verified from the manual.', space_after=8)

add_table(
    headers=['Category', 'Count', 'Notes'],
    rows=[
        ['Calib. Preset modes',                      '9',    'Cinema Film 1/2, Reference, TV, Photo, Game, Bright Cinema, Bright TV, User'],
        ['Core tone controls (Contrast/Brightness/Color/Hue/Sharpness)', '5', 'Basic level'],
        ['Image processing toggles (Reality Creation, Iris, Contrast Enhancer, Lamp, Motionflow)', '5', 'Each has multiple sub-options'],
        ['Colour Temperature presets',               '4 fixed + 5 custom = 9', 'D93, D75, D65, D55, Custom 1–5'],
        ['Colour Temperature Gain/Bias sub-controls','6 per slot × 9 slots = 54', 'Full expert colour temp control'],
        ['Gamma Correction presets',                 '10 + Off',   '6 standard + 4 custom LDT slots (Gamma 7–10)'],
        ['Colour Correction matrix',                 '18',   '6 colours × Hue, Saturation, Brightness'],
        ['Noise & processing (NR, MPEG NR, Smooth Gradation, Film Mode)', '4', 'Each multi-option'],
        ['Color Space — global gamut sliders',  '2',   'Cyan–Red and Magenta–Green; apply to the entire selected color space, not per primary. Verified from OSD screenshot'],
        ['Remaining Expert (Clear White, x.v.Color, HDR, Input Lag)', '4', ''],
        ['Advanced Picture (Auto Calibration)',      '4 functions', 'Pre Check, Adjust, Before/After, Reset'],
    ],
    col_widths=[3.1, 1.4, 1.9],
    note='Counts reflect distinct named controls. Sub-options and slider ranges within each control are additional degrees of freedom.'
)

add_callout(
    'The Colour Temperature Gain/Bias section alone (54 adjustments across 9 slots) exceeds the total '
    'count of all other basic controls combined — reflecting the depth of calibration available on this projector.',
    fill_hex='EEF4FF', border_hex='006DC6', label='KEY INSIGHT', label_color=ACCENT
)

# ─── Footer rule ───
add_rule('CCCCCC', 4)
p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_foot.paragraph_format.space_before = Pt(2)
rf = p_foot.add_run('Source: VPL-VW385ES Operating Instructions (Sony Corporation, 47254410M)  ·  Verified 2026')
rf.font.name  = 'Calibri'
rf.font.size  = Pt(8)
rf.font.color.rgb = MID_GRAY

# ─── Save ───
out = '/Users/johnlee/Downloads/sony-ldt-editor/Resources/VPL-VW385ES_Calibration_Controls.docx'
doc.save(out)
print(f'Saved: {out}')
