from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.page_width  = Inches(11)
section.page_height = Inches(8.5)
section.left_margin   = Inches(0.75)
section.right_margin  = Inches(0.75)
section.top_margin    = Inches(0.75)
section.bottom_margin = Inches(0.75)

# ─── Palette ───
BLACK       = RGBColor(0x1A, 0x1A, 0x1A)
DARK_GRAY   = RGBColor(0x44, 0x44, 0x44)
MID_GRAY    = RGBColor(0x77, 0x77, 0x77)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
SONY_BLUE   = RGBColor(0x00, 0x3B, 0x8E)
ACCENT      = RGBColor(0x00, 0x6D, 0xC6)
GREEN       = RGBColor(0x1A, 0x7A, 0x3C)
GREEN_BG    = RGBColor(0xE6, 0xF4, 0xEC)
AMBER       = RGBColor(0xB8, 0x6C, 0x00)
AMBER_BG    = RGBColor(0xFF, 0xF8, 0xEC)
GRAY_BG     = RGBColor(0xF2, 0xF2, 0xF2)
RED         = RGBColor(0xAA, 0x22, 0x22)

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'),   kwargs[edge].get('val','single'))
            tag.set(qn('w:sz'),    kwargs[edge].get('sz','4'))
            tag.set(qn('w:space'),'0')
            tag.set(qn('w:color'), kwargs[edge].get('color','auto'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def remove_table_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'none')
        tblBorders.append(tag)
    tblPr.append(tblBorders)

def add_heading(text, level=1, color=SONY_BLUE, size=None, space_before=14, space_after=4, bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = color
    sizes = {1: 16, 2: 12, 3: 10}
    run.font.size = Pt(size or sizes.get(level, 11))
    run.font.name = 'Calibri'
    return p

def add_body(text, color=DARK_GRAY, size=9.5, space_before=2, space_after=4):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name   = 'Calibri'
    run.font.size   = Pt(size)
    run.font.color.rgb = color
    return p

def add_rule(color_hex='CCCCCC', thickness=4):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(thickness))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    return p

# ── Cover ──────────────────────────────────────────────────────────────
p_cover = doc.add_paragraph()
p_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cover.paragraph_format.space_before = Pt(4)
p_cover.paragraph_format.space_after  = Pt(2)
tc = p_cover._p.get_or_add_pPr()
shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'003B8E'); tc.append(shd)
rc = p_cover.add_run('VPL-VW385ES  ·  Complete Menu Inventory')
rc.font.name='Calibri'; rc.font.size=Pt(18); rc.font.color.rgb=WHITE; rc.bold=True

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after  = Pt(6)
tc2 = p_sub._p.get_or_add_pPr()
shd2 = OxmlElement('w:shd'); shd2.set(qn('w:val'),'clear'); shd2.set(qn('w:color'),'auto'); shd2.set(qn('w:fill'),'003B8E'); tc2.append(shd2)
rs = p_sub.add_run('All menu items categorised by type  ·  Green highlight = already implemented in app')
rs.font.name='Calibri'; rs.font.size=Pt(10); rs.font.color.rgb=RGBColor(0xAA,0xC8,0xFF)

add_rule('1A6FD8', 8)

# ── Legend ─────────────────────────────────────────────────────────────
add_heading('Legend', level=2, color=DARK_GRAY, space_before=6, space_after=4)
legend_rows = [
    ('✅  In App',        'E6F4EC', '1A7A3C', 'Already implemented and working in the current app'),
    ('🎛  Adjustable',    'EEF4FF', '003B8E', 'Can be changed via SDCP SET command — user-facing control'),
    ('▶  Action',        'FFF8EC', 'B86C00', 'One-shot command (no persistent value) — e.g. Reset, Calibrate'),
    ('👁  Read-Only',     'F2F2F2', '444444', 'Status/information display only — no SET possible'),
    ('🔧  Installation',  'FFF0F0', 'AA2222', 'One-time physical install setting — set once and forget'),
]
tbl_legend = doc.add_table(rows=1+len(legend_rows), cols=3)
tbl_legend.alignment = WD_TABLE_ALIGNMENT.LEFT
remove_table_borders(tbl_legend)
hdr = tbl_legend.rows[0]
for i, h in enumerate(['Badge', 'Category', 'Meaning']):
    shade_cell(hdr.cells[i], '003B8E')
    p = hdr.cells[i].paragraphs[0]
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(h); r.bold=True; r.font.name='Calibri'; r.font.size=Pt(9); r.font.color.rgb=WHITE
for ri, (badge, bg, fc, desc) in enumerate(legend_rows):
    row = tbl_legend.rows[ri+1]
    shade_cell(row.cells[0], bg); shade_cell(row.cells[1], bg); shade_cell(row.cells[2], bg)
    for ci, txt in enumerate([badge, badge.split('  ')[1], desc]):
        p = row.cells[ci].paragraphs[0]
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(txt if ci != 0 else badge)
        r.font.name='Calibri'; r.font.size=Pt(9)
        r.font.color.rgb = RGBColor(int(fc[0:2],16), int(fc[2:4],16), int(fc[4:6],16))
        if ci == 0: r.bold = True
for row in tbl_legend.rows:
    row.cells[0].width = Inches(1.2)
    row.cells[1].width = Inches(1.3)
    row.cells[2].width = Inches(6.5)

add_rule()

# ── Main table ─────────────────────────────────────────────────────────
# Columns: #, Menu, Sub-section, Item Name, Type, Options/Range, Notes, In App
# in_app: True = green, False = white

ROWS = [
    # (#, menu, sub, item, type, options, notes, in_app)

    # ── PICTURE MENU ──
    (1,  'Picture', 'Core Tone',      'Calib. Preset',        '🎛 Adjustable', '9 modes: Cinema Film 1/2, Reference, TV, Photo, Game, Bright Cinema, Bright TV, User', 'Per-input, per-mode storage', True),
    (2,  'Picture', 'Core Tone',      'Reset',                '▶ Action',      '—',                                          'Resets current Calib. Preset to factory defaults', False),
    (3,  'Picture', 'Core Tone',      'Contrast',             '🎛 Adjustable', '0–100',                                      'Label changes to Contrast(HDR) when HDR10/HLG active', True),
    (4,  'Picture', 'Core Tone',      'Brightness',           '🎛 Adjustable', '0–100',                                      '',  True),
    (5,  'Picture', 'Core Tone',      'Color',                '🎛 Adjustable', 'Slider',                                     'Master saturation — all hues equally', False),
    (6,  'Picture', 'Core Tone',      'Hue',                  '🎛 Adjustable', 'Slider',                                     'Master hue rotation (–red / +green)', False),
    (7,  'Picture', 'Core Tone',      'Color Temp',           '🎛 Adjustable', 'D93 / D75 / D65 / D55 / Custom 1–5',        '',  True),
    (8,  'Picture', 'Color Temp',     'Custom Gain R/G/B',    '🎛 Adjustable', '3 sliders per custom slot',                  '5 custom slots × 3 = 15 Gain sub-controls. Item codes TBD.', False),
    (9,  'Picture', 'Color Temp',     'Custom Bias R/G/B',    '🎛 Adjustable', '3 sliders per custom slot',                  '5 custom slots × 3 = 15 Bias sub-controls. Item codes TBD.', False),
    (10, 'Picture', 'Core Tone',      'Sharpness',            '🎛 Adjustable', 'Slider',                                     'Remote shortcut: SHARPNESS +/–', False),
    (11, 'Picture', 'Cinema Black Pro','Advanced Iris',        '🎛 Adjustable', 'Off / Full / Limited',                       'Dynamic iris control', True),
    (12, 'Picture', 'Cinema Black Pro','Contrast Enhancer',   '🎛 Adjustable', 'Off / Low / Middle / High',                  'Auto bright/dark balance per scene', False),
    (13, 'Picture', 'Cinema Black Pro','Lamp Control',        '🎛 Adjustable', 'High / Low',                                 'Low = deeper blacks, quieter, longer lamp life', False),
    (14, 'Picture', 'Core',           'Reality Creation',     '🎛 Adjustable', 'On / Off + sub-controls',                    'Sub: Database (Normal/4K), Resolution slider, Noise Filtering slider', False),
    (15, 'Picture', 'Core',           'Motionflow',           '🎛 Adjustable', 'Off / Smooth High / Smooth Low / Impulse / Combination / True Cinema', '', True),

    # ── EXPERT SETTING ──
    (16, 'Expert Setting', 'Noise',   'NR (Noise Reduction)', '🎛 Adjustable', 'Off / Low / Medium / High / Auto',           '',  True),
    (17, 'Expert Setting', 'Noise',   'MPEG NR',              '🎛 Adjustable', 'Off / Low / Medium / High / Auto',           'Block noise + mosquito noise from digital compression', False),
    (18, 'Expert Setting', 'Noise',   'Smooth Gradation',     '🎛 Adjustable', 'Off / Low / Medium / High',                  'Reduces banding in flat gradient areas', False),
    (19, 'Expert Setting', 'Noise',   'Film Mode',            '🎛 Adjustable', 'Auto / Off',                                 'Progressive detection for film-sourced content', False),
    (20, 'Expert Setting', 'Gamma',   'Gamma Correction',     '🎛 Adjustable', '1.8 / 2.0 / 2.1 / 2.2 / 2.4 / 2.6 / Gamma 7–10 / Off', 'Gamma 7–10 = custom LDT slots', True),
    (21, 'Expert Setting', 'Color Correction', 'Red: Hue',       '🎛 Adjustable', '−50 to +50',  'Item 00h 87h', False),
    (22, 'Expert Setting', 'Color Correction', 'Red: Saturation','🎛 Adjustable', '−50 to +50',  'Item 00h 88h', False),
    (23, 'Expert Setting', 'Color Correction', 'Red: Brightness','🎛 Adjustable', '−30 to +30',  'Item 00h 89h', False),
    (24, 'Expert Setting', 'Color Correction', 'Yellow: Hue',    '🎛 Adjustable', '−50 to +50',  'Item 00h 8Ah', False),
    (25, 'Expert Setting', 'Color Correction', 'Yellow: Saturation','🎛 Adjustable','−50 to +50', 'Item 00h 8Bh', False),
    (26, 'Expert Setting', 'Color Correction', 'Yellow: Brightness','🎛 Adjustable','−30 to +30', 'Item 00h 8Ch', False),
    (27, 'Expert Setting', 'Color Correction', 'Green: Hue',     '🎛 Adjustable', '−50 to +50',  'Item 00h 8Dh', False),
    (28, 'Expert Setting', 'Color Correction', 'Green: Saturation','🎛 Adjustable','−50 to +50',  'Item 00h 8Eh', False),
    (29, 'Expert Setting', 'Color Correction', 'Green: Brightness','🎛 Adjustable','−30 to +30',  'Item 00h 8Fh', False),
    (30, 'Expert Setting', 'Color Correction', 'Cyan: Hue',      '🎛 Adjustable', '−50 to +50',  'Item 00h 90h', False),
    (31, 'Expert Setting', 'Color Correction', 'Cyan: Saturation','🎛 Adjustable', '−50 to +50',  'Item 00h 91h', False),
    (32, 'Expert Setting', 'Color Correction', 'Cyan: Brightness','🎛 Adjustable', '−30 to +30',  'Item 00h 92h', False),
    (33, 'Expert Setting', 'Color Correction', 'Blue: Hue',      '🎛 Adjustable', '−50 to +50',  'Item 00h 93h', False),
    (34, 'Expert Setting', 'Color Correction', 'Blue: Saturation','🎛 Adjustable', '−50 to +50',  'Item 00h 94h', False),
    (35, 'Expert Setting', 'Color Correction', 'Blue: Brightness','🎛 Adjustable', '−30 to +30',  'Item 00h 95h', False),
    (36, 'Expert Setting', 'Color Correction', 'Magenta: Hue',      '🎛 Adjustable','−50 to +50', 'Item 00h 96h', False),
    (37, 'Expert Setting', 'Color Correction', 'Magenta: Saturation','🎛 Adjustable','−50 to +50', 'Item 00h 97h', False),
    (38, 'Expert Setting', 'Color Correction', 'Magenta: Brightness','🎛 Adjustable','−30 to +30', 'Item 00h 98h', False),
    (39, 'Expert Setting', 'Color Space', 'Color Space',        '🎛 Adjustable', 'BT.709 / BT.2020 / CS1 / CS2 / CS3 / Custom', 'Item 00h 3Bh', True),
    (40, 'Expert Setting', 'Color Space', 'Custom: Cyan–Red',   '🎛 Adjustable', '−100 to +100', 'Shown only when Custom selected. Item 00h 76h', True),
    (41, 'Expert Setting', 'Color Space', 'Custom: Magenta–Green','🎛 Adjustable','−100 to +100', 'Shown only when Custom selected. Item 00h 77h', True),
    (42, 'Expert Setting', 'Other',    'Clear White',           '🎛 Adjustable', 'Off / Low / High',                          'Emphasises vivid whites. Item code TBD.', False),
    (43, 'Expert Setting', 'Other',    'x.v.Color',             '🎛 Adjustable', 'On / Off',                                  'Disables gamma correction when On. Item code TBD.', False),
    (44, 'Expert Setting', 'Other',    'HDR',                   '🎛 Adjustable', 'Off / HDR10 / HLG / Auto',                  'Item 00h 7Ch', True),
    (45, 'Expert Setting', 'Other',    'Input Lag Reduction',   '🎛 Adjustable', 'On / Off',                                  'When On: Motionflow/NR/MPEG NR disabled. Item 00h 99h', False),

    # ── ADVANCED PICTURE ──
    (46, 'Advanced Picture', 'Auto Calibration', 'Pre Check',   '▶ Action', '—', 'Measures panel color vs factory reference; reports dE. Takes minutes.', False),
    (47, 'Advanced Picture', 'Auto Calibration', 'Adjust',      '▶ Action', '—', 'Performs automatic panel color calibration. Lens moves. Takes minutes.', False),
    (48, 'Advanced Picture', 'Auto Calibration', 'Before/After','▶ Action', '—', 'Toggles factory vs calibrated at fixed frequency for visual comparison', False),
    (49, 'Advanced Picture', 'Auto Calibration', 'Reset',       '▶ Action', '—', 'Reverts panel calibration to factory defaults', False),

    # ── SCREEN MENU ──
    (50, 'Screen', '', 'Aspect',      '🎛 Adjustable', 'Normal / 1.85:1 Zoom / 2.35:1 Zoom / V Stretch / Squeeze / Stretch', 'Item code TBD.', False),
    (51, 'Screen', '', 'Blanking: Left',   '🎛 Adjustable', 'Slider', 'Crop left edge. Item code TBD.', False),
    (52, 'Screen', '', 'Blanking: Right',  '🎛 Adjustable', 'Slider', 'Crop right edge. Item code TBD.', False),
    (53, 'Screen', '', 'Blanking: Top',    '🎛 Adjustable', 'Slider', 'Crop top edge. Item code TBD.', False),
    (54, 'Screen', '', 'Blanking: Bottom', '🎛 Adjustable', 'Slider', 'Crop bottom edge. Item code TBD.', False),
    (55, 'Screen', '', 'Picture Position','🔧 Installation','Save/Delete/Select (up to 5 presets)', 'Stores lens position + aspect + blanking. Triggers lens motor movement.', False),

    # ── FUNCTION MENU ──
    (56, 'Function', '3D Settings', '2D-3D Display',         '🎛 Adjustable', 'Auto / 3D / 2D',                 'Item code TBD.', False),
    (57, 'Function', '3D Settings', '3D Format',             '🎛 Adjustable', 'Simulated 3D / Side-by-Side / Over-Under', 'Item code TBD.', False),
    (58, 'Function', '3D Settings', '3D Brightness',         '🎛 Adjustable', 'High / Standard',                'Item code TBD.', False),
    (59, 'Function', '3D Settings', '3D Depth Adjust',       '🎛 Adjustable', '−2 / −1 / 0 / +1 / +2',         'Item code TBD.', False),
    (60, 'Function', '3D Settings', 'Simulated 3D Effect',   '🎛 Adjustable', 'High / Middle / Low',            'Item code TBD.', False),
    (61, 'Function', '',            'Dynamic Range',          '🎛 Adjustable', 'Auto / Limited / Full',          'Sets HDMI input level. Item code TBD.', False),
    (62, 'Function', '',            'HDMI Signal Format',     '🎛 Adjustable', 'Standard / Enhanced',            '4K HDMI format. Item code TBD.', False),
    (63, 'Function', '',            'Test Pattern',           '🎛 Adjustable', 'On / Off',                       'Green test pattern for lens focus/zoom/shift. Item code TBD.', False),
    (64, 'Function', '',            'Settings Lock',          '🎛 Adjustable', 'Off / Level A / Level B',        'Locks menu items to prevent accidental changes. Item code TBD.', False),

    # ── SETUP MENU ──
    (65, 'Setup', '', 'OSD Status',          '🔧 Installation', 'On / Off',        'Turns on-screen display on/off. Item code TBD.', False),
    (66, 'Setup', '', 'Language',            '🔧 Installation', 'Multiple',        'Menu display language. Item code TBD.', False),
    (67, 'Setup', '', 'Menu Position',       '🔧 Installation', 'Bottom Left / Center', 'Item code TBD.', False),
    (68, 'Setup', '', 'High Altitude Mode',  '🔧 Installation', 'On / Off',        'Required above 1,500 m. Item code TBD.', False),
    (69, 'Setup', '', 'Remote Start',        '🔧 Installation', 'On / Off',        'Network wake-on-LAN. Item code TBD.', False),
    (70, 'Setup', '', 'Network Management',  '🔧 Installation', 'On / Off',        'Keep network active in standby. Item code TBD.', False),
    (71, 'Setup', '', 'Power Saving',        '🔧 Installation', 'Off / Standby',   'Auto standby after 10 min no signal. Item code TBD.', False),
    (72, 'Setup', '', 'All Reset',           '▶ Action',        '—',               'Factory reset all settings (not Lamp Timer). Item code TBD.', False),

    # ── INSTALLATION MENU ──
    (73, 'Installation', '', 'Image Flip',        '🔧 Installation', 'Off / HV / H / V',           'Ceiling/rear projection orientation. Item code TBD.', False),
    (74, 'Installation', '', 'Lens Control',       '🔧 Installation', 'On / Off',                    'Locks lens motor buttons. Item code TBD.', False),
    (75, 'Installation', '', 'Anamorphic Lens',    '🔧 Installation', '1.24x / 1.32x',               'Anamorphic conversion ratio. Item code TBD.', False),
    (76, 'Installation', '', 'Trigger Select',     '🔧 Installation', 'Off / Power / V Stretch / 2.35:1 Zoom', '12V trigger output. Item code TBD.', False),
    (77, 'Installation', '', 'IR Receiver',        '🔧 Installation', 'Front & Rear / Front / Rear', 'Remote control detector location. Item code TBD.', False),
    (78, 'Installation', '', 'Panel Alignment',    '🔧 Installation', 'Shift/Zone per R/B vs G',     'Color convergence correction per zone. Complex procedure. Item code TBD.', False),
    (79, 'Installation', '', 'Network Setting',    '🔧 Installation', 'DHCP / Manual IP',            'IP address configuration. Managed via projector OSD or web UI.', False),

    # ── INFORMATION MENU ──
    (80, 'Information', '', 'Model Name',      '👁 Read-Only', '—', '', False),
    (81, 'Information', '', 'Serial Number',   '👁 Read-Only', '—', '', False),
    (82, 'Information', '', 'Signal Type',     '👁 Read-Only', '—', 'Resolution + 2D/3D + frame format', False),
    (83, 'Information', '', 'Color Format',    '👁 Read-Only', '—', 'YCbCr / RGB', False),
    (84, 'Information', '', 'Color Space',     '👁 Read-Only', '—', 'BT.601 / BT.709 / BT.2020', False),
    (85, 'Information', '', 'HDR Format',      '👁 Read-Only', '—', 'SDR / HDR(HDR10) / HDR(HLG)', False),
    (86, 'Information', '', 'Software Version','👁 Read-Only', '—', '', False),
    (87, 'Information', '', 'Lamp Timer',      '👁 Read-Only', '—', 'Total hours of lamp usage', False),

    # ── POWER / INPUT (SDCP, not OSD) ──
    (88, 'Power/Input', '', 'Power On',       '▶ Action',     '—',                    'SDCP command. Item 01h 30h value 01h', False),
    (89, 'Power/Input', '', 'Power Off',      '▶ Action',     '—',                    'SDCP command. Item 01h 30h value 00h', False),
    (90, 'Power/Input', '', 'Input Select',   '🎛 Adjustable','HDMI 1 / HDMI 2',      'Item code TBD.', False),
]

# Category color mapping
TYPE_COLORS = {
    '✅': ('E6F4EC', '1A7A3C'),
    '🎛': ('EEF4FF', '003B8E'),
    '▶':  ('FFF8EC', 'B86C00'),
    '👁': ('F2F2F2', '444444'),
    '🔧': ('FFF0F0', 'AA2222'),
}

def get_type_colors(type_str, in_app):
    if in_app:
        return ('D4EDDA', '1A7A3C')
    emoji = type_str[0]
    return TYPE_COLORS.get(emoji, ('FFFFFF', '000000'))

# ── Build section tables ────────────────────────────────────────────────
HEADERS = ['#', 'Menu', 'Sub-Section', 'Item Name', 'Type', 'Options / Range', 'Notes']
COL_W   = [0.35, 1.1, 1.3, 1.7, 1.2, 2.1, 2.0]  # total ~9.75

current_menu = None

# Group rows by menu
from itertools import groupby
def menu_key(r): return r[1]

# Add table per menu section
menu_groups = {}
for row in ROWS:
    m = row[1]
    if m not in menu_groups:
        menu_groups[m] = []
    menu_groups[m].append(row)

for menu_name, rows in menu_groups.items():
    add_heading(menu_name, level=1, space_before=10, space_after=3)

    tbl = doc.add_table(rows=1 + len(rows), cols=len(HEADERS))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(tbl)

    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(HEADERS):
        shade_cell(hdr_row.cells[i], '003B8E')
        set_cell_border(hdr_row.cells[i], bottom={'val':'single','sz':'4','color':'FFFFFF'})
        p = hdr_row.cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(h.upper())
        r.bold=True; r.font.name='Calibri'; r.font.size=Pt(8.5); r.font.color.rgb=WHITE

    # Data rows
    for ri, row_data in enumerate(rows):
        num, menu, sub, item, rtype, options, notes, in_app = row_data
        bg, fc_hex = get_type_colors(rtype, in_app)
        fc = RGBColor(int(fc_hex[0:2],16), int(fc_hex[2:4],16), int(fc_hex[4:6],16))

        cells = tbl.rows[ri+1].cells
        values = [str(num), menu, sub, item, ('✅ ' if in_app else '') + rtype, options, notes]

        for ci, (cell, val) in enumerate(zip(cells, values)):
            shade_cell(cell, bg)
            set_cell_border(cell, bottom={'val':'single','sz':'2','color':'D0D0D0'})
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(val)
            r.font.name = 'Calibri'
            r.font.size = Pt(8.5)
            if ci == 3:  # Item name bold
                r.bold = True
                r.font.color.rgb = RGBColor(int('1A',16),int('1A',16),int('1A',16))
            elif in_app and ci == 4:
                r.bold = True
                r.font.color.rgb = GREEN
            else:
                r.font.color.rgb = RGBColor(int('44',16),int('44',16),int('44',16))

    # Set column widths
    for row in tbl.rows:
        for ci, w in enumerate(COL_W):
            row.cells[ci].width = Inches(w)

# ── Footer ──────────────────────────────────────────────────────────────
add_rule()
p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_foot.paragraph_format.space_before = Pt(2)
rf = p_foot.add_run('Source: VPL-VW385ES Operating Instructions (Sony, 47254410M)  ·  Generated 2026-04-16  ·  90 items total')
rf.font.name='Calibri'; rf.font.size=Pt(8); rf.font.color.rgb=MID_GRAY

out = '/Users/johnlee/Downloads/sony-ldt-editor/Resources/VPL-VW385ES_Menu_Inventory.docx'
doc.save(out)
print(f'Saved: {out}')
