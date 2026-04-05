"""
Generates two .docx reports:
1. HDR/SDR Calibration Slot Persistence
2. Three Calibration Scenarios Planning Guide
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────
#  SHARED HELPERS
# ─────────────────────────────────────────────────────────────
def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.page_width    = Inches(8.5)
    s.page_height   = Inches(11)
    s.left_margin   = Inches(1.0)
    s.right_margin  = Inches(1.0)
    s.top_margin    = Inches(0.85)
    s.bottom_margin = Inches(0.85)
    return doc

# Palette
BLACK      = RGBColor(0x1A,0x1A,0x1A)
DARK_GRAY  = RGBColor(0x44,0x44,0x44)
MID_GRAY   = RGBColor(0x77,0x77,0x77)
WHITE      = RGBColor(0xFF,0xFF,0xFF)
SONY_BLUE  = RGBColor(0x00,0x3B,0x8E)
ACCENT     = RGBColor(0x00,0x6D,0xC6)
AMBER      = RGBColor(0xB8,0x6C,0x00)
RED_DARK   = RGBColor(0x99,0x1A,0x1A)
GREEN_DARK = RGBColor(0x1A,0x6B,0x2E)
PURPLE     = RGBColor(0x5B,0x2D,0x8E)

# Scenario accent colors
C_SDR709   = RGBColor(0x00,0x55,0xA0)   # blue
C_HDR      = RGBColor(0xB8,0x6C,0x00)   # amber
C_P3       = RGBColor(0x1A,0x6B,0x2E)   # green

HEX_SDR709 = '0055A0'
HEX_HDR    = 'B86C00'
HEX_P3     = '1A6B2E'
HEX_BLUE   = '003B8E'

def shade_para(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), hex_color); pPr.append(shd)

def shade_cell(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

def border_cell(cell, **kw):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        if edge in kw:
            t = OxmlElement(f'w:{edge}')
            t.set(qn('w:val'),  kw[edge].get('val','single'))
            t.set(qn('w:sz'),   kw[edge].get('sz','4'))
            t.set(qn('w:space'),'0')
            t.set(qn('w:color'),kw[edge].get('color','auto'))
            b.append(t)
    tcPr.append(b)

def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    tb = OxmlElement('w:tblBorders')
    for e in ('top','left','bottom','right','insideH','insideV'):
        t = OxmlElement(f'w:{e}'); t.set(qn('w:val'),'none'); tb.append(t)
    tblPr.append(tb)

def hrule(doc, color='CCCCCC', sz=4, sb=4, sa=8):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),str(sz))
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),color)
    pBdr.append(bot); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)

def cover_banner(doc, title, subtitle, fill='003B8E', sub_color='AAC8FF'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    shade_para(p, fill)
    r = p.add_run(title)
    r.font.name='Calibri'; r.font.size=Pt(18)
    r.font.color.rgb=WHITE; r.bold=True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(4)
    shade_para(p2, fill)
    r2 = p2.add_run(subtitle)
    r2.font.name='Calibri'; r2.font.size=Pt(10)
    r2.font.color.rgb=RGBColor(int(sub_color[0:2],16),int(sub_color[2:4],16),int(sub_color[4:6],16))
    hrule(doc, '1A6FD8', 8, sb=0, sa=10)

def heading(doc, text, level=1, color=SONY_BLUE, size=None,
            sb=14, sa=6, bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(text)
    r.bold=bold; r.font.name='Calibri'
    r.font.color.rgb=color
    r.font.size=Pt(size or {1:18,2:13,3:11,4:10}.get(level,10))
    return p

def body(doc, text, color=DARK_GRAY, size=10.5, italic=False, sb=2, sa=6, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(sb)
    p.paragraph_format.space_after=Pt(sa)
    r = p.add_run(text)
    r.font.name='Calibri'; r.font.size=Pt(size)
    r.font.color.rgb=color; r.italic=italic; r.bold=bold
    return p

def callout(doc, text, fill='EEF4FF', border='006DC6',
            label=None, lcolor=None, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(10)
    p.paragraph_format.left_indent=Inches(0.15)
    p.paragraph_format.right_indent=Inches(0.15)
    shade_para(p, fill)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    for edge in ('top','left','bottom','right'):
        t = OxmlElement(f'w:{edge}')
        t.set(qn('w:val'),'single')
        t.set(qn('w:sz'),'12' if edge=='left' else '4')
        t.set(qn('w:space'),'4'); t.set(qn('w:color'),border); pBdr.append(t)
    pPr.append(pBdr)
    if label:
        lc = lcolor or ACCENT
        rl = p.add_run(label+'   ')
        rl.bold=True; rl.font.name='Calibri'
        rl.font.size=Pt(size); rl.font.color.rgb=lc
    rb = p.add_run(text)
    rb.font.name='Calibri'; rb.font.size=Pt(size); rb.font.color.rgb=DARK_GRAY

def mono_block(doc, lines, fill='F4F4F4'):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
        p.paragraph_format.left_indent=Inches(0.5)
        shade_para(p, fill)
        r = p.add_run(line)
        r.font.name='Courier New'; r.font.size=Pt(9.5); r.font.color.rgb=DARK_GRAY
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before=Pt(0); sp.paragraph_format.space_after=Pt(8)

def section_banner(doc, text, fill_hex, text_color=WHITE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(8)
    shade_para(p, fill_hex)
    r = p.add_run('  ' + text)
    r.bold=True; r.font.name='Calibri'; r.font.size=Pt(13)
    r.font.color.rgb=text_color

def data_table(doc, headers, rows, col_widths=None, hdr_fill='003B8E', alt_fill='F7F9FF'):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    no_borders(t)
    # header
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade_cell(c, hdr_fill)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(5)
        r = p.add_run(h.upper())
        r.bold=True; r.font.name='Calibri'
        r.font.size=Pt(8.5); r.font.color.rgb=WHITE
    # rows
    for ri, row_data in enumerate(rows):
        tr = t.rows[ri+1]
        fill = alt_fill if ri%2==1 else 'FFFFFF'
        for ci, txt in enumerate(row_data):
            c = tr.cells[ci]
            shade_cell(c, fill)
            c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            border_cell(c, bottom={'val':'single','sz':'2','color':'E0E0E0'})
            p = c.paragraphs[0]
            p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)
            r = p.add_run(txt)
            r.font.name='Calibri'; r.font.size=Pt(9.5)
            r.font.color.rgb=BLACK; r.bold=(ci==0)
    if col_widths:
        for i,w in enumerate(col_widths):
            for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)
    return t

def footer(doc, text):
    hrule(doc,'CCCCCC',4,sb=14,sa=4)
    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name='Calibri'; r.font.size=Pt(8); r.font.color.rgb=MID_GRAY

# ═════════════════════════════════════════════════════════════
#  REPORT 1 — HDR / SDR CALIBRATION SLOT PERSISTENCE
# ═════════════════════════════════════════════════════════════
def build_report1():
    doc = new_doc()

    cover_banner(doc,
        'HDR / SDR Calibration Slot Persistence',
        'VPL-VW385ES  ·  How Gamma Slots and Settings Are Stored by Signal Type  ·  Verified from Operating Instructions (47254410M)')

    # ── 1. The Core Question
    heading(doc,'1  The Core Question',1,sb=6)
    body(doc,
        'When you upload a custom gamma curve (.ldt file) to one of the Gamma 7–10 slots via '
        'Sony ImageDirector, does that curve survive when you switch between an SDR signal '
        '(Rec.709) and an HDR signal (Rec.2020 / HDR10)? And does the projector maintain '
        'separate settings for each signal type?')

    callout(doc,
        'The uploaded curve in any Gamma slot (7–10) is stored permanently in that slot. '
        'Switching signal types does NOT erase or overwrite it. The slot holds the curve '
        'until you deliberately overwrite it with another upload.',
        fill='EEF4FF', border='006DC6', label='KEY FINDING', lcolor=ACCENT)

    # ── 2. How Memory Is Organized
    heading(doc,'2  How the Projector Organizes Its Memory',1)
    body(doc,
        'The projector uses a two-level storage system. Understanding this is essential for '
        'planning where to save calibration settings for each signal type.')

    heading(doc,'2.1  Level 1 — Preset Memory Numbers (Signal-Driven)',2,color=ACCENT,sb=10,sa=4)
    body(doc,
        'The projector automatically assigns an internal preset memory number to every incoming '
        'signal based on its resolution and frame rate. This happens automatically — you cannot '
        'choose or name these numbers. They are the projector\'s way of recognizing signal types.')

    data_table(doc,
        headers=['Preset Numbers','Resolution / Signal','Typical Use'],
        rows=[
            ['5–14',    '480p, 576p, 720p, 1080i, 1080p (HD)',      'SDR Blu-ray, streaming SDR, broadcast'],
            ['74–79',   '3840×2160 & 4096×2160 (4K UHD, 50/60Hz)', 'HDR10 streaming, 4K sources'],
            ['93–96',   '3840×2160 & 4096×2160 (4K UHD, 24/25/30Hz)','HDR10 Blu-ray, cinema-rate 4K'],
            ['18–20, 71–72', '3D signals',                           'Active 3D content'],
            ['26–55',   'Computer/PC signals',                       'PC input'],
        ],
        col_widths=[1.4,2.8,2.2],hdr_fill='003B8E')

    callout(doc,
        'Your SDR 1080p Blu-ray player triggers preset numbers 5–14. '
        'Your HDR10 4K source triggers preset numbers 74–79 or 93–96. '
        'These are different memory banks — they do not share settings.',
        fill='FFFBEA', border='B86C00', label='PRACTICAL MEANING', lcolor=AMBER)

    heading(doc,'2.2  Level 2 — Calib. Preset (User-Driven)',2,color=ACCENT,sb=10,sa=4)
    body(doc,
        'Within each preset memory number, the projector stores settings per Calib. Preset mode. '
        'The manual confirms: Gamma Correction, Color Space, Color Temp, CMS, and HDR setting '
        'are all stored "for each input connector and each Calib. Preset." '
        'This gives you 9 independent setting profiles per input.')

    data_table(doc,
        headers=['What Is Stored Per','Storage Scope'],
        rows=[
            ['Gamma Correction (slot selection)','Per input connector + per Calib. Preset'],
            ['Color Space',                      'Per input connector + per Calib. Preset'],
            ['Color Temp (Gain/Bias R/G/B)',      'Per Color Temp slot (D93–D55, Custom 1–5)'],
            ['Color Correction CMS',              'Per input connector + per Calib. Preset'],
            ['HDR setting',                       'Per input connector + per Calib. Preset'],
            ['Calib. Preset selection',           'Per input connector'],
            ['Gamma slot curve data (7–10)',       'Global — one copy, shared across all presets'],
        ],
        col_widths=[3.0,3.4],hdr_fill='003B8E')

    callout(doc,
        'The gamma curve data stored in each slot (7–10) is GLOBAL — it is shared across '
        'all Calib. Presets and all inputs. Only the SELECTION of which slot to use is stored '
        'per preset. This means: if you overwrite Gamma 7 with a new curve, every preset '
        'that was using Gamma 7 will now use the new curve.',
        fill='FFF0F0', border='991A1A', label='CRITICAL DISTINCTION', lcolor=RED_DARK)

    # ── 3. Visual Map
    heading(doc,'3  Memory Map — What Lives Where',1)
    mono_block(doc,[
        '  GAMMA SLOT DATA (Global — permanent until overwritten)',
        '  ┌──────────────┬──────────────┬──────────────┬──────────────┐',
        '  │  Gamma 7     │  Gamma 8     │  Gamma 9     │  Gamma 10    │',
        '  │  (your curve)│  (your curve)│  (your curve)│  (spare)     │',
        '  └──────────────┴──────────────┴──────────────┴──────────────┘',
        '',
        '  HDMI 1 INPUT — HD signal (preset 5–14) = SDR source',
        '  ├─ Calib. Preset: Reference → Gamma: Slot 7, CS: BT.709, HDR: Off',
        '  ├─ Calib. Preset: Cinema Film 2 → Gamma: Slot 9, CS: BT.2020, HDR: Off',
        '  └─ (other presets independent)',
        '',
        '  HDMI 1 INPUT — 4K signal (preset 74–96) = HDR source',
        '  ├─ Calib. Preset: Cinema Film 1 → Gamma: Slot 8, CS: BT.2020, HDR: HDR10',
        '  └─ (other presets independent)',
    ])

    # ── 4. What Happens When You Switch
    heading(doc,'4  What Happens When You Switch Signal Types',1)

    heading(doc,'4.1  Different Sources on Different HDMI Ports',2,color=ACCENT,sb=10,sa=4)
    body(doc,'This is the ideal setup. Each port maintains its own independent settings.')
    data_table(doc,
        headers=['HDMI Port','Source','Preset Numbers','Recalls'],
        rows=[
            ['HDMI 1','4K HDR media player','74–96','Cinema Film 1 → Gamma 8, BT.2020, HDR10'],
            ['HDMI 2','SDR 1080p Blu-ray',  '5–14', 'Reference → Gamma 7, BT.709, HDR Off'],
        ],
        col_widths=[1.0,1.8,1.4,2.2],hdr_fill='1A6B2E')
    callout(doc,
        'Switching inputs automatically loads the correct Calib. Preset settings for that source. '
        'No manual switching needed — the ideal configuration.',
        fill='F0FFF4', border='1A6B2E', label='BEST PRACTICE', lcolor=GREEN_DARK)

    heading(doc,'4.2  SDR and HDR from the Same Device on One Port',2,color=ACCENT,sb=10,sa=4)
    body(doc,
        'If a single device (e.g. Apple TV 4K, media player) outputs both SDR and HDR '
        'content on the same HDMI port, the projector may still detect different preset '
        'memory numbers if the resolution changes (e.g. SDR at 1080p vs HDR at 4K). '
        'However, if both are 4K resolution, they share the same preset memory number.')

    data_table(doc,
        headers=['Scenario','Preset Number','Setting Recall','Action Required'],
        rows=[
            ['SDR source outputs 1080p, HDR outputs 4K','Different (5–14 vs 74–96)',
             'Different Calib. Presets load automatically','None — works automatically'],
            ['Both SDR and HDR output at 4K from same device','Same (74–96)',
             'Same Calib. Preset is recalled for both','Manually switch Calib. Preset when changing content type'],
        ],
        col_widths=[2.2,1.3,1.6,1.3],hdr_fill='B86C00')

    callout(doc,
        'When HDR Auto is set to On, the projector reads HDR metadata from the signal and '
        'can override your Color Space setting (forces BT.2020 for HDR10 signals). '
        'This makes calibration unpredictable. Set HDR manually in each calibrated Calib. Preset.',
        fill='FFF0F0', border='991A1A', label='WARNING — HDR AUTO', lcolor=RED_DARK)

    # ── 5. Summary
    heading(doc,'5  Quick Reference Summary',1)
    data_table(doc,
        headers=['Question','Answer'],
        rows=[
            ['Does the uploaded curve survive a signal change?',
             'Yes — Gamma slot data is global and permanent until overwritten'],
            ['Are there separate HDR and SDR gamma slots?',
             'No — the 4 slots (7–10) are shared. You assign which slot each Calib. Preset uses'],
            ['Are settings stored separately for HD vs 4K signals?',
             'Yes — different preset memory numbers (5–14 for HD, 74–96 for 4K) each have independent Calib. Preset storage'],
            ['Does switching Color Space to BT.2020 erase Gamma slot data?',
             'No — Color Space and Gamma Correction are independent settings'],
            ['If I overwrite Gamma 7 with a new curve, does it affect other presets?',
             'Yes — slot data is global. Any preset pointing to Gamma 7 will use the new curve'],
        ],
        col_widths=[3.0,3.4],hdr_fill='003B8E')

    footer(doc,'Sony VPL-VW385ES  ·  HDR/SDR Calibration Slot Persistence  ·  Verified from Operating Instructions (47254410M)')

    out = '/Users/johnlee/Downloads/sony-ldt-editor/Resources/Report1_Calibration_Slot_Persistence.docx'
    doc.save(out)
    print(f'Saved: {out}')


# ═════════════════════════════════════════════════════════════
#  REPORT 2 — THREE CALIBRATION SCENARIOS
# ═════════════════════════════════════════════════════════════
def build_report2():
    doc = new_doc()

    cover_banner(doc,
        'VPL-VW385ES — Three Calibration Scenarios',
        'Complete Planning Guide  ·  SDR Rec.709  ·  HDR Rec.2020 PQ  ·  SDR DCI-P3  ·  Sony VPL-VW385ES')

    # ── Overview
    heading(doc,'Overview',1,sb=6)
    body(doc,
        'This guide covers three distinct calibration targets for the Sony VPL-VW385ES. '
        'Each scenario requires a different combination of Color Space, Gamma Correction slot, '
        'HDR setting, and Calib. Preset. The goal is to store each calibration independently '
        'so they can be recalled without interference.')

    # Scenario summary cards (3-col table)
    t = doc.add_table(rows=2, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    no_borders(t)
    cards = [
        ('SCENARIO 1','SDR · Rec.709\nGamma 2.4 / BT.1886','Calib. Preset: Reference\nGamma Slot: 7\nColor Space: BT.709\nHDR: Off', HEX_SDR709),
        ('SCENARIO 2','HDR · Rec.2020\nPQ Curve','Calib. Preset: Cinema Film 1\nGamma Slot: 8\nColor Space: BT.2020\nHDR: HDR10', HEX_HDR),
        ('SCENARIO 3','SDR · DCI-P3\nGamma 2.4 / BT.1886','Calib. Preset: Cinema Film 2\nGamma Slot: 9\nColor Space: BT.2020 + CMS\nHDR: Off', HEX_P3),
    ]
    for ci,(title,sub,detail,hexcol) in enumerate(cards):
        # header row
        c0 = t.rows[0].cells[ci]
        shade_cell(c0, hexcol)
        border_cell(c0,top={'val':'single','sz':'8','color':hexcol})
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_before=Pt(6); p0.paragraph_format.space_after=Pt(2)
        r0 = p0.add_run(title)
        r0.bold=True; r0.font.name='Calibri'; r0.font.size=Pt(11); r0.font.color.rgb=WHITE
        p0b = c0.add_paragraph()
        p0b.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0b.paragraph_format.space_before=Pt(0); p0b.paragraph_format.space_after=Pt(6)
        r0b = p0b.add_run(sub)
        r0b.font.name='Calibri'; r0b.font.size=Pt(9.5); r0b.font.color.rgb=WHITE
        # detail row
        c1 = t.rows[1].cells[ci]
        shade_cell(c1, 'F8F8F8')
        border_cell(c1,
            bottom={'val':'single','sz':'8','color':hexcol},
            left={'val':'single','sz':'4','color':'E0E0E0'},
            right={'val':'single','sz':'4','color':'E0E0E0'})
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before=Pt(6); p1.paragraph_format.space_after=Pt(6)
        p1.paragraph_format.left_indent=Inches(0.1)
        r1 = p1.add_run(detail)
        r1.font.name='Calibri'; r1.font.size=Pt(9); r1.font.color.rgb=DARK_GRAY
        c0.width=Inches(2.13); c1.width=Inches(2.13)

    doc.add_paragraph().paragraph_format.space_after=Pt(8)

    # ── Color Space Menu Clarification
    heading(doc,'1  Color Space Menu — What Exists on This Projector',1)
    body(doc,
        'Before planning the scenarios, it\'s important to know exactly what Color Space '
        'options the projector provides — and what is notably absent.')

    data_table(doc,
        headers=['Menu Option','Type','Gamut Coverage','Use For'],
        rows=[
            ['BT.709',       'Fixed standard', 'Rec.709 / sRGB',        'Scenario 1 — SDR broadcast and Blu-ray'],
            ['BT.2020',      'Fixed standard', 'Full UHD wide gamut',   'Scenario 2 & 3 — HDR and P3 (starting point)'],
            ['Color Space 1','Sony preset',    'TV/sports tuning',      'General TV use — not for calibration'],
            ['Color Space 2','Sony preset',    'Bright room TV tuning', 'General TV use — not for calibration'],
            ['Color Space 3','Sony preset',    'Bright room cinema',    'General use — not for calibration'],
            ['Custom',       'User-adjustable','Global shift only (2 sliders)','Minor gamut trim only — not P3 capable'],
            ['DCI-P3',       '⚠ NOT AVAILABLE','—',                    'Must be achieved via BT.2020 + CMS corrections'],
        ],
        col_widths=[1.3,1.2,1.6,2.3],hdr_fill='003B8E')

    callout(doc,
        'There is no DCI-P3 color space preset on the VW385ES. P3 calibration requires '
        'selecting BT.2020 as the color space and using the Color Correction CMS to pull '
        'each primary and secondary to P3 target coordinates.',
        fill='FFF0F0', border='991A1A', label='IMPORTANT', lcolor=RED_DARK)

    # ── Gamma Slot Plan
    heading(doc,'2  Gamma Slot Assignment',1)
    body(doc,
        'The four custom gamma slots (7–10) hold uploaded .ldt curves. Assign one slot '
        'per scenario to keep them fully independent. The built-in gamma options '
        '(1.8 / 2.0 / 2.1 / 2.2 / 2.4 / 2.6) do not occupy these slots.')

    data_table(doc,
        headers=['Slot','Assigned Scenario','Curve Type','Generated With'],
        rows=[
            ['Gamma 7','Scenario 1 — SDR Rec.709',  'BT.1886, Lb ≈ 0.5%',        'LDT Editor → BT.1886 preset, Lb slider at 0.5%'],
            ['Gamma 8','Scenario 2 — HDR PQ',        'PQ tone map, ~200 nit',      'LDT Editor → PQ preset, nit slider at 200'],
            ['Gamma 9','Scenario 3 — SDR DCI-P3',    'BT.1886, Lb ≈ 0.5%',        'Same curve as Gamma 7 — separate slot for independence'],
            ['Gamma 10','(Spare)',                    '—',                          'Reserve for future use or experimentation'],
        ],
        col_widths=[0.8,1.9,1.7,2.0],hdr_fill='5B2D8E')

    callout(doc,
        'Scenarios 1 and 3 use the same BT.1886 curve but stored in separate slots (7 and 9). '
        'This ensures that adjusting or re-uploading the curve for one scenario never '
        'accidentally affects the other.',
        fill='EEF4FF', border='006DC6', label='WHY TWO SLOTS FOR THE SAME CURVE', lcolor=ACCENT)

    # ── SCENARIO 1
    section_banner(doc,'SCENARIO 1  —  SDR · Rec.709 · Gamma 2.4 (BT.1886)', HEX_SDR709)

    data_table(doc,
        headers=['Setting','Value','Reason'],
        rows=[
            ['Calib. Preset',    'Reference',        'Designed for faithful, unprocessed source reproduction — ideal for SDR calibration baseline'],
            ['HDR',              'Off  (manual)',    'Hard-set Off — prevents Auto mode from overriding Color Space when a BT.2020-flagged signal appears'],
            ['Color Space',      'BT.709',           'Correct gamut container for Rec.709 SDR content'],
            ['Gamma Correction', 'Gamma 7',          'BT.1886 curve uploaded via LDT Editor. More accurate than built-in 2.4 — compensates for projector\'s non-zero black level'],
            ['Color Temp',       'Custom 3  (D65)',  'D65 is the Rec.709 white point. Custom 3 factory default matches D65'],
            ['Color Correction', 'CMS to Rec.709 targets','After setting Gamma and Color Temp, measure primaries/secondaries and dial in CMS corrections'],
            ['Lamp Control',     'Low (recommended)','Low lamp deepens blacks, improves contrast for dark-room SDR viewing'],
        ],
        col_widths=[1.5,1.6,3.3],hdr_fill=HEX_SDR709)

    callout(doc,
        'Why BT.1886 instead of the built-in 2.4? The built-in gamma 2.4 is a pure power law '
        'that assumes a perfect black level of zero. The VW385ES has a very low but non-zero '
        'black level. BT.1886 accounts for this with a soft toe near black, giving more accurate '
        'shadow rendering that matches how reference monitors handle Rec.709 content.',
        fill='EEF4FF', border='0055A0', label='BT.1886 vs BUILT-IN 2.4', lcolor=C_SDR709)

    # ── SCENARIO 2
    section_banner(doc,'SCENARIO 2  —  HDR · Rec.2020 · PQ Curve', HEX_HDR)

    data_table(doc,
        headers=['Setting','Value','Reason'],
        rows=[
            ['Calib. Preset',    'Cinema Film 1',    'Cinema-oriented preset, natural home for HDR mastered theatrical content'],
            ['HDR',              'HDR10  (manual)',  'Set manually to HDR10 — do NOT use Auto. Auto can override Color Space mid-calibration and causes unpredictable behavior'],
            ['Color Space',      'BT.2020',          'HDR10 content is delivered in the BT.2020 container'],
            ['Gamma Correction', 'Gamma 8',          'PQ tone-mapping curve uploaded via LDT Editor, targeting ~200 nit (VW385ES peak brightness)'],
            ['Color Temp',       'Custom 3  (D65)',  'HDR10 mastering white point is D65'],
            ['Color Correction', 'CMS to P3-D65 targets','HDR10 content is typically mastered within a P3-D65 gamut inside the BT.2020 container — target P3 primaries, not full BT.2020'],
            ['Lamp Control',     'High',             'HDR requires maximum brightness headroom for specular highlights'],
        ],
        col_widths=[1.5,1.6,3.3],hdr_fill=HEX_HDR)

    callout(doc,
        'Calibrate CMS to P3-D65 primary coordinates, not full BT.2020 primaries. '
        'The vast majority of HDR10 content is graded within a P3 gamut and delivered '
        'in a BT.2020 container. Calibrating to full BT.2020 primaries will result in '
        'oversaturated colors for most real-world HDR content.',
        fill='FFFBEA', border='B86C00', label='CMS TARGET FOR HDR', lcolor=AMBER)

    mono_block(doc,[
        '  PQ Curve Generation (LDT Editor):',
        '  1. Open LDT Editor → Sidebar → HDR PQ section',
        '  2. Set nit slider to 200 (VW385ES peak brightness)',
        '  3. Click "Generate PQ → 200 nit"',
        '  4. Export As → name file "HDR_PQ_200nit.ldt"',
        '  5. Open Sony ImageDirector → load file → upload to Gamma 8 slot',
    ])

    # ── SCENARIO 3
    section_banner(doc,'SCENARIO 3  —  SDR · DCI-P3 · Gamma 2.4 (BT.1886)', HEX_P3)

    data_table(doc,
        headers=['Setting','Value','Reason'],
        rows=[
            ['Calib. Preset',    'Cinema Film 2',    'Cinema Film 2 is appropriate for theatrical P3 content — distinct from Reference (Rec.709)'],
            ['HDR',              'Off  (manual)',    'SDR content — hard-set Off'],
            ['Color Space',      'BT.2020',          'No P3 preset exists. Start from BT.2020 (widest gamut) and trim to P3 targets via CMS. Starting from BT.709 would clip P3 primaries before CMS can correct them'],
            ['Gamma Correction', 'Gamma 9',          'Same BT.1886 curve as Scenario 1 but in a separate slot for independence'],
            ['Color Temp',       'Custom 3  (D65)',  'P3-D65 is the standard for home/streaming P3 content. Use D63 only for theatrical DCI alignment'],
            ['Color Correction', 'CMS to P3-D65 targets','This is how P3 is achieved — measure each primary/secondary and correct to P3 chromaticity coordinates'],
            ['Lamp Control',     'Low (recommended)','Dark-room SDR viewing'],
        ],
        col_widths=[1.5,1.6,3.3],hdr_fill=HEX_P3)

    callout(doc,
        'Why BT.2020 as starting point for P3?\n'
        'P3 primaries lie OUTSIDE Rec.709 but INSIDE BT.2020. '
        'If you start from BT.709, the projector clips any color that exceeds Rec.709 '
        'before the CMS can correct it — you can never reach P3. '
        'Starting from BT.2020 leaves the full P3 gamut accessible for CMS to shape.',
        fill='F0FFF4', border='1A6B2E', label='WHY BT.2020 FOR P3', lcolor=GREEN_DARK)

    mono_block(doc,[
        '  DCI-P3 Primary Targets (D65 white point):',
        '  Red:     x=0.680  y=0.320',
        '  Green:   x=0.265  y=0.690',
        '  Blue:    x=0.150  y=0.060',
        '  White:   x=0.3127 y=0.3290  (D65)',
        '',
        '  Use these as CMS targets when measuring with a colorimeter.',
        '  Rec.709 Red (x=0.640) and Green (x=0.300) are notably inside P3 —',
        '  those two primaries will require the most CMS correction.',
    ])

    # ── Complete Settings Comparison
    heading(doc,'3  Side-by-Side Settings Comparison',1)
    data_table(doc,
        headers=['Setting','Scenario 1\nSDR Rec.709','Scenario 2\nHDR Rec.2020','Scenario 3\nSDR DCI-P3'],
        rows=[
            ['Calib. Preset', 'Reference',       'Cinema Film 1', 'Cinema Film 2'],
            ['HDR',           'Off',              'HDR10',         'Off'],
            ['Color Space',   'BT.709',           'BT.2020',       'BT.2020'],
            ['Gamma Slot',    'Gamma 7',          'Gamma 8',       'Gamma 9'],
            ['Gamma Curve',   'BT.1886, Lb 0.5%', 'PQ 200 nit',   'BT.1886, Lb 0.5%'],
            ['Color Temp',    'Custom 3 (D65)',   'Custom 3 (D65)','Custom 3 (D65)'],
            ['CMS Targets',   'Rec.709 primaries','P3-D65 primaries','P3-D65 primaries'],
            ['Lamp Control',  'Low',              'High',          'Low'],
        ],
        col_widths=[1.5,1.7,1.7,1.5],hdr_fill='003B8E')

    # ── Calibration Order
    heading(doc,'4  Calibration Procedure Order  (Same for All Scenarios)',1)
    body(doc,'Always calibrate in this sequence. Each step depends on the previous being locked first.',sb=2,sa=6)

    t2 = doc.add_table(rows=6, cols=3)
    t2.alignment = WD_TABLE_ALIGNMENT.LEFT
    no_borders(t2)
    steps = [
        ('1','Color Space\n+ Cyan–Red / Magenta–Green',
         'Foundation — sets where the gamut triangle sits. Lock this first. Any change here invalidates all steps below.'),
        ('2','Color Temp\nGain R/G/B + Bias R/G/B',
         'Sets white point and black point within the gamut. Measure with colorimeter and adjust Custom Color Temp sub-controls.'),
        ('3','Gamma Correction\n(upload .ldt curve)',
         'Sets tonal response. Upload appropriate .ldt via ImageDirector. Verify with EOTF measurement.'),
        ('4','Color Correction (CMS)\nHue / Saturation / Brightness per color',
         'Fine-tune individual primaries and secondaries. Measure each color at 100% saturation and correct to target.'),
        ('5','Verify saturation sweep\n25% / 50% / 75% / 100%',
         'CMS corrections at 100% do not translate perfectly to lower saturations. Measure all levels and accept residual error within ΔE < 3.'),
        ('6','Save Calib. Preset',
         'Navigate to Calib. Preset → confirm all settings are correct → the projector saves automatically. Verify by recalling the preset.'),
    ]
    fills = ['E8F0FB','FFFFFF','E8F0FB','FFFFFF','E8F0FB','FFFFFF']
    for ri,(step,ctrl,why) in enumerate(steps):
        for ci,txt in enumerate([step,ctrl,why]):
            c = t2.rows[ri].cells[ci]
            shade_cell(c,fills[ri])
            border_cell(c,bottom={'val':'single','sz':'2','color':'D0D0D0'})
            if ri==0:
                border_cell(c,top={'val':'single','sz':'4','color':'003B8E'})
            p = c.paragraphs[0]
            p.paragraph_format.space_before=Pt(5)
            p.paragraph_format.space_after=Pt(5)
            if ci==0: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(txt)
            r.font.name='Calibri'; r.font.size=Pt(9.5); r.font.color.rgb=BLACK
            r.bold=(ci==0)
            if ci==1: r.font.color.rgb=SONY_BLUE; r.bold=True
        t2.rows[ri].cells[0].width=Inches(0.4)
        t2.rows[ri].cells[1].width=Inches(2.1)
        t2.rows[ri].cells[2].width=Inches(3.9)

    doc.add_paragraph().paragraph_format.space_after=Pt(6)

    callout(doc,
        'If you ever go back and adjust the Color Space or Cyan–Red/Magenta–Green slider '
        '(Step 1), you must redo Steps 2 through 6 entirely. Steps 3–6 are always downstream '
        'of the gamut foundation.',
        fill='FFF0F0',border='991A1A',label='CALIBRATION ORDER RULE',lcolor=RED_DARK)

    # ── Switching Workflow
    heading(doc,'5  Day-to-Day Switching Workflow',1)

    heading(doc,'If HDR and SDR sources are on different HDMI ports:',2,color=ACCENT,sb=8,sa=4,size=10)
    mono_block(doc,[
        '  Switch HDMI input  →  Projector auto-loads the correct Calib. Preset  →  Done.',
        '  No manual action required.',
    ],'F0FFF4')

    heading(doc,'If both come from one device on one HDMI port:',2,color=AMBER,sb=8,sa=4,size=10)
    mono_block(doc,[
        '  Content changes to HDR  →  Press CALIBRATED PRESET → select Cinema Film 1',
        '  Content changes to SDR Rec.709  →  Press CALIBRATED PRESET → select Reference',
        '  Content changes to SDR P3  →  Press CALIBRATED PRESET → select Cinema Film 2',
        '',
        '  The CALIBRATED PRESET button on the remote control makes this a single button press.',
    ],'FFFBEA')

    footer(doc,'Sony VPL-VW385ES  ·  Three Calibration Scenarios Planning Guide  ·  Based on Operating Instructions (47254410M) and OSD verification')

    out = '/Users/johnlee/Downloads/sony-ldt-editor/Resources/Report2_Three_Calibration_Scenarios.docx'
    doc.save(out)
    print(f'Saved: {out}')


# ─── Run both ───
build_report1()
build_report2()
print('Done.')
